import os
import sys
from pathlib import Path
import sqlite3
from tkinter import *
from tkinter import ttk, filedialog, messagebox
from datetime import datetime, date
from docx import Document
import matplotlib.pyplot as plt
from PIL import Image, ImageTk
import calendar
from tkcalendar import DateEntry
import webbrowser
from matplotlib.figure import Figure
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
import docx
from docx.shared import Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
# =============================================
# CONFIGURAÇÕES INICIAIS
# =============================================

COR_FUNDO = "#f8f9fa"
COR_PRIMARIA = "#2c3e50"
COR_SECUNDARIA = "#34495e"
COR_DESTAQUE = "#3498db"
COR_TEXTO = "#2c3e50"
COR_CARD = "#ffffff"
COR_ALERTA = "#e74c3c"
COR_SUCESSO = "#2ecc71"

# Configurar diretórios
def get_app_data_dir():
    """Retorna o diretório correto para cada sistema operacional"""
    home = Path.home()
    
    if sys.platform == "win32":
        return home / "AppData" / "Local" / "SeuApp"
    elif sys.platform == "linux":
        return home / ".local" / "share" / "seuapp"  # Padrão no Linux
    elif sys.platform == "darwin":  # MacOS
        return home / "Library" / "Application Support" / "SeuApp"
    else:
        return Path(__file__).parent  # Fallback



APP_DATA_DIR = Path.home() / "AppData" / "Local" / "SeuApp"
COMPROVANTES_DIR = APP_DATA_DIR / "comprovantes"
RELATORIOS_DIR = APP_DATA_DIR / "relatorios"


try:
    COMPROVANTES_DIR.mkdir(parents=True, exist_ok=True)
    RELATORIOS_DIR.mkdir(parents=True, exist_ok=True)
except Exception as e:
    print(f"Erro ao criar pastas: {e}")

# =============================================
# BANCO DE DADOS - ESTRUTURA SIMPLIFICADA
# =============================================

def inicializar_banco_dados():
    conn = sqlite3.connect("sistema.db")
    cursor = conn.cursor()
    
    # Tabela de clientes (proprietários)
    cursor.execute("""
    CREATE TABLE IF NOT EXISTS clientes (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        nome TEXT NOT NULL,
        telefone TEXT,
        email TEXT,
        endereco TEXT
    );
    """)
    
    # Tabela de imóveis
    cursor.execute("""
    CREATE TABLE IF NOT EXISTS imoveis (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        cliente_id INTEGER,
        endereco TEXT NOT NULL,
        quartos INTEGER,
        banheiros INTEGER,
        plataforma TEXT,  
        FOREIGN KEY (cliente_id) REFERENCES clientes (id)
    );
    """)
        
    # Tabela de serviços de limpeza
    cursor.execute("""
    CREATE TABLE IF NOT EXISTS limpezas (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        imovel_id INTEGER,
        data DATE NOT NULL,
        hora_inicio TEXT,
        hora_fim TEXT,
        horas_trabalhadas REAL,
        valor_hora REAL DEFAULT 30.0,
        valor_total REAL,
        observacoes TEXT,
        FOREIGN KEY (imovel_id) REFERENCES imoveis (id)
    );
    """)
    
    # Tabela de tipos de enxoval 
    cursor.execute("""
    CREATE TABLE IF NOT EXISTS tipos_enxoval (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        nome TEXT NOT NULL,
        preco_unitario REAL DEFAULT 0.0,
        unidade_medida TEXT DEFAULT 'unidade',
        data_cadastro DATE DEFAULT CURRENT_DATE
    )
    """)
    
    # Tabela de consumo de enxoval
    cursor.execute("""
    CREATE TABLE IF NOT EXISTS consumo_enxoval (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        imovel_id INTEGER NOT NULL,
        item_id INTEGER NOT NULL,
        quantidade INTEGER NOT NULL,
        data DATE NOT NULL,
        FOREIGN KEY (imovel_id) REFERENCES imoveis(id),
        FOREIGN KEY (item_id) REFERENCES tipos_enxoval(id)
    )
    """)
    
    # Tabela de suprimentos
    cursor.execute("""
    CREATE TABLE IF NOT EXISTS suprimentos (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        nome TEXT NOT NULL,
        preco_unitario REAL DEFAULT 0.0,
        unidade_medida TEXT DEFAULT 'unidade',
        data_cadastro DATE DEFAULT CURRENT_DATE
    )
    """)
    
    # Tabela de reposição de suprimentos
    cursor.execute("""
    CREATE TABLE IF NOT EXISTS reposicao_suprimentos (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        imovel_id INTEGER,
        suprimento_id INTEGER,
        data DATE NOT NULL,
        quantidade INTEGER DEFAULT 1,
        valor_gasto REAL,
        comprovante_path TEXT,
        FOREIGN KEY (imovel_id) REFERENCES imoveis (id),
        FOREIGN KEY (suprimento_id) REFERENCES suprimentos (id)
    );
    """)


    cursor.execute("""
        CREATE TABLE IF NOT EXISTS fechamentos (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            tipo TEXT NOT NULL,  -- 'imovel' ou 'cliente'
            referencia_id INTEGER NOT NULL,  -- id do imóvel ou cliente
            data_inicio DATE NOT NULL,
            data_fim DATE NOT NULL,
            data_fechamento DATE DEFAULT CURRENT_DATE,
            valor_total REAL NOT NULL,
            observacoes TEXT,
            comprovante_path TEXT
    );  
    """)

    # Inserir dados básicos se as tabelas estiverem vazias
    cursor.execute("SELECT COUNT(*) FROM tipos_enxoval")
    if cursor.fetchone()[0] == 0:
        itens_enxoval = [
            ('Lençol Solteiro', 45.0, 'unidade'),
            ('Lençol Casal', 65.0, 'unidade'),
            ('Lençol Queen', 75.0, 'unidade'),
            ('Lençol King', 85.0, 'unidade'),
            ('Toalha de Banho', 35.0, 'unidade'),
            ('Toalha de Rosto', 25.0, 'unidade'),
            ('Toalha de Mesa', 30.0, 'unidade'),
            ('Cobertor', 90.0, 'unidade'),
            ('Edredom', 120.0, 'unidade')
        ]
        cursor.executemany("INSERT INTO tipos_enxoval (nome, preco_unitario, unidade_medida) VALUES (?, ?, ?)", itens_enxoval)
    
    cursor.execute("SELECT COUNT(*) FROM suprimentos")
    if cursor.fetchone()[0] == 0:
        suprimentos = [
            ('Sabonete', 1.5, 'unidade'),
            ('Shampoo', 5.0, 'unidade'),
            ('Condicionador', 5.0, 'unidade'),
            ('Papel Higiênico', 0.5, 'rolo'),
            ('Papel Toalha', 2.0, 'rolo'),
            ('Café', 15.0, 'pacote'),
            ('Açúcar', 8.0, 'kg'),
            ('Sabão em Pó', 12.0, 'kg'),
            ('Amaciante', 10.0, 'litro')
        ]
        cursor.executemany("INSERT INTO suprimentos (nome, preco_unitario, unidade_medida) VALUES (?, ?, ?)", suprimentos)
    

    conn.commit()
    conn.close()

# =============================================
# FUNÇÕES AUXILIARES
# =============================================

def calcular_horas(hora_inicio, hora_fim):
    """Calcula a diferença entre duas horas no formato HH:MM"""
    try:
        inicio = datetime.strptime(hora_inicio, "%H:%M")
        fim = datetime.strptime(hora_fim, "%H:%M")
        diferenca = fim - inicio
        horas = diferenca.seconds / 3600
        return round(horas, 2)
    except ValueError:
        return 0.0

def selecionar_arquivo(entry_widget, diretorio="comprovantes"):
    """Permite selecionar um arquivo e move para o diretório de comprovantes"""
    caminho = filedialog.askopenfilename()
    if caminho:
        nome_arquivo = os.path.basename(caminho)
        destino = os.path.join(diretorio, nome_arquivo)
        os.replace(caminho, destino)
        entry_widget.delete(0, END)
        entry_widget.insert(0, destino)

def formatar_moeda(valor):
    """Formata um valor float para string monetária"""
    return f"R$ {valor:.2f}".replace(".", ",")

def gerar_relatorio_semanal():
    """Gera um relatório semanal em formato DOCX"""
    conn = sqlite3.connect("sistema.db")
    cursor = conn.cursor()
    
    # Data de início (7 dias atrás)
    data_inicio = (datetime.now() - timedelta(days=7)).strftime("%Y-%m-%d")
    
    # Criar documento
    doc = Document()
    doc.add_heading('Relatório Semanal de Serviços', 0)
    
    # Adicionar data
    doc.add_paragraph(f"Período: {data_inicio} a {datetime.now().strftime('%Y-%m-%d')}")
    
    # Limpezas realizadas
    doc.add_heading('Limpezas Realizadas', level=1)
    cursor.execute("""
        SELECT i.endereco, l.data, l.horas_trabalhadas, l.valor_total 
        FROM limpezas l
        JOIN imoveis i ON l.imovel_id = i.id
        WHERE date(l.data) >= date(?)
        ORDER BY l.data
    """, (data_inicio,))
    
    tabela_limpezas = doc.add_table(rows=1, cols=4)
    tabela_limpezas.style = 'Light Shading'
    cabecalho = tabela_limpezas.rows[0].cells
    cabecalho[0].text = 'Imóvel'
    cabecalho[1].text = 'Data'
    cabecalho[2].text = 'Horas'
    cabecalho[3].text = 'Valor'
    
    total_horas = 0
    total_valor = 0
    
    for row in cursor.fetchall():
        linha = tabela_limpezas.add_row().cells
        linha[0].text = row[0]
        linha[1].text = row[1]
        linha[2].text = str(row[2])
        linha[3].text = formatar_moeda(row[3])
        total_horas += row[2]
        total_valor += row[3]
    
    doc.add_paragraph(f"Total de horas trabalhadas: {total_horas}")
    doc.add_paragraph(f"Total a receber por limpezas: {formatar_moeda(total_valor)}")
    
    # Enxoval utilizado
    doc.add_heading('Enxoval Utilizado', level=1)
    cursor.execute("""
        SELECT i.endereco, t.nome, SUM(c.quantidade), t.preco_unitario
        FROM consumo_enxoval c
        JOIN imoveis i ON c.imovel_id = i.id
        JOIN tipos_enxoval t ON c.item_id = t.id
        WHERE date(c.data) >= date(?)
        GROUP BY i.endereco, t.nome
        ORDER BY i.endereco
    """, (data_inicio,))
    
    tabela_enxoval = doc.add_table(rows=1, cols=4)
    tabela_enxoval.style = 'Light Shading'
    cabecalho = tabela_enxoval.rows[0].cells
    cabecalho[0].text = 'Imóvel'
    cabecalho[1].text = 'Item'
    cabecalho[2].text = 'Quantidade'
    cabecalho[3].text = 'Valor Unitário'
    
    total_itens = 0
    total_valor_enxoval = 0
    
    for row in cursor.fetchall():
        linha = tabela_enxoval.add_row().cells
        linha[0].text = row[0]
        linha[1].text = row[1]
        linha[2].text = str(row[2])
        linha[3].text = formatar_moeda(row[3])
        total_itens += row[2]
        total_valor_enxoval += row[2] * row[3]
    
    doc.add_paragraph(f"Total de itens utilizados: {total_itens}")
    doc.add_paragraph(f"Total a receber por enxoval: {formatar_moeda(total_valor_enxoval)}")
    
    # Suprimentos repostos
    doc.add_heading('Suprimentos Repostos', level=1)
    cursor.execute("""
        SELECT i.endereco, s.nome, SUM(r.quantidade), r.valor_gasto
        FROM reposicao_suprimentos r
        JOIN imoveis i ON r.imovel_id = i.id
        JOIN suprimentos s ON r.suprimento_id = s.id
        WHERE date(r.data) >= date(?)
        GROUP BY i.endereco, s.nome
        ORDER BY i.endereco
    """, (data_inicio,))
    
    tabela_suprimentos = doc.add_table(rows=1, cols=4)
    tabela_suprimentos.style = 'Light Shading'
    cabecalho = tabela_suprimentos.rows[0].cells
    cabecalho[0].text = 'Imóvel'
    cabecalho[1].text = 'Item'
    cabecalho[2].text = 'Quantidade'
    cabecalho[3].text = 'Valor Gasto'
    
    total_suprimentos = 0
    total_valor_suprimentos = 0
    
    for row in cursor.fetchall():
        linha = tabela_suprimentos.add_row().cells
        linha[0].text = row[0]
        linha[1].text = row[1]
        linha[2].text = str(row[2])
        linha[3].text = formatar_moeda(row[3])
        total_suprimentos += row[2]
        total_valor_suprimentos += row[3]
    
    doc.add_paragraph(f"Total de suprimentos repostos: {total_suprimentos}")
    doc.add_paragraph(f"Total gasto com suprimentos: {formatar_moeda(total_valor_suprimentos)}")
    doc.add_paragraph(f"Valor fixo semanal por gestão: {formatar_moeda(50.0 * cursor.execute('SELECT COUNT(DISTINCT imovel_id) FROM limpezas WHERE date(data) >= date(?)', (data_inicio,)).fetchone()[0])}")
    
    # Total geral
    doc.add_heading('Resumo Financeiro', level=1)
    total_geral = total_valor + total_valor_enxoval + (50.0 * cursor.execute('SELECT COUNT(DISTINCT imovel_id) FROM limpezas WHERE date(data) >= date(?)', (data_inicio,)).fetchone()[0])
    doc.add_paragraph(f"Total a receber: {formatar_moeda(total_geral)}", style='Heading 2')
    
    # Salvar documento
    nome_arquivo = f"relatorios/Relatorio_{datetime.now().strftime('%Y%m%d')}.docx"
    doc.save(nome_arquivo)
    conn.close()
    
    return nome_arquivo

# =============================================
# INTERFACE PRINCIPAL
# =============================================

class SistemaGestaoApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Sistema de Gestão de Serviços para Airbnb")
        self.root.geometry("1100x700")
        self.root.configure(bg=COR_FUNDO)
        
        # Configurações de estilo
        self.fonte_titulo = ("Helvetica", 14, "bold")
        self.fonte_normal = ("Helvetica", 11)
        self.fonte_pequena = ("Helvetica", 9)
        
        # Inicializar banco de dados
        inicializar_banco_dados()
        
        # Criar layout principal
        self.criar_cabecalho()
        self.criar_menu_lateral()
        self.criar_area_conteudo()
        
        # Mostrar tela inicial
        self.mostrar_tela("dashboard")
    
    def criar_cabecalho(self):
        """Cria o cabeçalho do sistema"""
        frame = Frame(self.root, bg=COR_PRIMARIA, height=70)
        frame.pack(fill=X)
        
        Label(frame, 
              text="GESTÃO DE SERVIÇOS ", 
              bg=COR_PRIMARIA, 
              fg="white", 
              font=("Helvetica", 16, "bold")).pack(side=LEFT, padx=20)
        
        self.label_data = Label(frame, 
                              text=datetime.now().strftime("%d/%m/%Y %H:%M"),
                              bg=COR_PRIMARIA,
                              fg="white",
                              font=self.fonte_normal)
        self.label_data.pack(side=RIGHT, padx=20)
    
    def criar_menu_lateral(self):
        """Cria o menu de navegação lateral"""
        self.frame_menu = Frame(self.root, bg=COR_SECUNDARIA, width=200)
        self.frame_menu.pack(side=LEFT, fill=Y)
        
        # Menu de módulos
        modulos = [
            ("Dashboard", "dashboard"),
            ("Clientes", "clientes"),
            ("Imóveis", "imoveis"),
            ("Limpeza", "limpeza"),
            ("Enxoval", "enxoval"),
            ("Configurar Itens", "config_itens"),
            ("Suprimentos", "suprimentos"),
            ("Relatórios", "relatorios"),
            ("Fechar Contas", "fechar_contas")
        ]
        
        for texto, comando in modulos:
            btn = Button(self.frame_menu,
                        text=texto,
                        command=lambda c=comando: self.mostrar_tela(c),
                        bg=COR_SECUNDARIA,
                        fg="white",
                        activebackground=COR_DESTAQUE,
                        activeforeground="white",
                        borderwidth=0,
                        font=self.fonte_normal,
                        anchor="w",
                        padx=20,
                        pady=12)
            btn.pack(fill=X)
    
    def criar_area_conteudo(self):
        """Cria a área de conteúdo principal"""
        self.frame_conteudo = Frame(self.root, bg=COR_FUNDO)
        self.frame_conteudo.pack(fill=BOTH, expand=True, padx=20, pady=20)
        
        # Criar todos os frames de conteúdo
        self.criar_dashboard()
        self.criar_clientes()
        self.criar_imoveis()
        self.criar_limpeza()
        self.criar_enxoval()
        self.criar_suprimentos()
        self.criar_relatorios()
        self.criar_config_itens()
        self.criar_fechamento_contas()
   


    def mostrar_tela(self, tela):
        """Controla qual tela mostrar"""
        for widget in self.frame_conteudo.winfo_children():
            widget.pack_forget()
        
        if tela == "dashboard":
            self.frame_dashboard.pack(fill=BOTH, expand=True)
            self.atualizar_dashboard()
        elif tela == "clientes":
            self.frame_clientes.pack(fill=BOTH, expand=True)
            self.carregar_clientes()
        elif tela == "imoveis":
            self.frame_imoveis.pack(fill=BOTH, expand=True)
            self.carregar_imoveis()
        elif tela == "limpeza":
            self.frame_limpeza.pack(fill=BOTH, expand=True)
            self.carregar_limpezas()
        elif tela == "enxoval":
            self.frame_enxoval.pack(fill=BOTH, expand=True)
            self.carregar_itens_enxoval()
        elif tela == "suprimentos":
            self.frame_suprimentos.pack(fill=BOTH, expand=True)
            self.carregar_suprimentos()
        elif tela == "relatorios":
            self.frame_relatorios.pack(fill=BOTH, expand=True)
        elif tela == "config_itens":
            self.frame_config_itens.pack(fill=BOTH, expand=True)
        elif tela == "fechar_contas":
            self.frame_fechamento.pack(fill=BOTH, expand=True)

    # =============================================
    # MÓDULO DASHBOARD
    # =============================================

    def criar_dashboard(self):
        """Cria o dashboard inicial"""
        self.frame_dashboard = Frame(self.frame_conteudo, bg=COR_FUNDO)

        # Frame para os cards de resumo
        frame_cards = Frame(self.frame_dashboard, bg=COR_FUNDO)
        frame_cards.pack(fill=X, pady=10)

        # Cards de resumo
        self.card_limpezas = self.criar_card(frame_cards, "Limpezas Esta Semana", "0", COR_DESTAQUE)
        self.card_enxoval = self.criar_card(frame_cards, "Enxoval Utilizado", "0", COR_SECUNDARIA)
        self.card_suprimentos = self.criar_card(frame_cards, "Suprimentos Repostos", "0", COR_PRIMARIA)
        self.card_receber = self.criar_card(frame_cards, "Total a Receber", "R$ 0,00", COR_SUCESSO)

        # Frame para os gráficos
        frame_graficos = Frame(self.frame_dashboard, bg=COR_FUNDO)
        frame_graficos.pack(fill=BOTH, expand=True)

        # ============================
        # Gráfico de Limpezas por Dia
        # ============================
        frame_grafico1 = Frame(frame_graficos, bg=COR_FUNDO)
        frame_grafico1.pack(side=LEFT, fill=BOTH, expand=True, padx=5, pady=5)
        Label(frame_grafico1, text="Limpezas por Dia", bg=COR_FUNDO, font=self.fonte_titulo).pack()

        fig1 = Figure(figsize=(4, 2.5), dpi=100)
        ax1 = fig1.add_subplot(111)
        dias = ['Seg', 'Ter', 'Qua', 'Qui', 'Sex']
        quantidades = [3, 5, 2, 4, 6]
        ax1.bar(dias, quantidades, color='#4CAF50')  # verde moderno
        ax1.set_facecolor('#FFFFFF')
        fig1.patch.set_facecolor(COR_CARD)
        ax1.tick_params(labelsize=8)
        ax1.spines['top'].set_visible(False)
        ax1.spines['right'].set_visible(False)

        self.canvas_grafico1 = FigureCanvasTkAgg(fig1, master=frame_grafico1)
        self.canvas_grafico1.draw()
        self.canvas_grafico1.get_tk_widget().pack(fill=BOTH, expand=True)

        # ==========================================
        # Gráfico de Itens de Enxoval Mais Utilizados
        # ==========================================
        frame_grafico2 = Frame(frame_graficos, bg=COR_FUNDO)
        frame_grafico2.pack(side=LEFT, fill=BOTH, expand=True, padx=5, pady=5)
        Label(frame_grafico2, text="Itens de Enxoval Mais Utilizados", bg=COR_FUNDO, font=self.fonte_titulo).pack()

        fig2 = Figure(figsize=(4, 2.5), dpi=100)
        ax2 = fig2.add_subplot(111)
        itens = ['Toalhas', 'Lençóis', 'Fronhas', 'Cobertores']
        usos = [40, 30, 20, 10]
        cores = ['#2196F3', '#FFC107', '#FF5722', '#9C27B0']
        ax2.pie(usos, labels=itens, autopct='%1.1f%%', colors=cores, startangle=140)
        fig2.patch.set_facecolor(COR_CARD)

        self.canvas_grafico2 = FigureCanvasTkAgg(fig2, master=frame_grafico2)
        self.canvas_grafico2.draw()
        self.canvas_grafico2.get_tk_widget().pack(fill=BOTH, expand=True)

    
    def criar_card(self, parent, titulo, valor, cor):
        """Cria um card de resumo para o dashboard"""
        card = Frame(parent, bg=COR_CARD, bd=0, 
                    highlightthickness=1, highlightbackground="#e0e0e0")
        card.pack(side=LEFT, padx=10, fill=BOTH, expand=True)
        
        Label(card, text=titulo, bg=COR_CARD, fg=COR_TEXTO, 
              font=self.fonte_titulo).pack(pady=(10, 5), padx=10, anchor="w")
        
        label_valor = Label(card, text=valor, bg=COR_CARD, fg=cor, 
                          font=("Helvetica", 24, "bold"))
        label_valor.pack(pady=(0, 10))
        
        return label_valor
    
    def atualizar_dashboard(self):
        """Atualiza os dados do dashboard"""
        conn = sqlite3.connect("sistema.db")
        cursor = conn.cursor()
        
        # Limpezas esta semana
        cursor.execute("""
            SELECT COUNT(*), SUM(valor_total) 
            FROM limpezas 
            WHERE date(data) >= date('now', '-7 days')
        """)
        limpezas = cursor.fetchone()
        self.card_limpezas.config(text=f"{limpezas[0]}\n{formatar_moeda(limpezas[1] or 0)}")
        
        # Enxoval utilizado
        cursor.execute("""
            SELECT COUNT(*), SUM(t.preco_unitario * c.quantidade)
            FROM consumo_enxoval c
            JOIN tipos_enxoval t ON c.item_id = t.id
            WHERE date(c.data) >= date('now', '-7 days')
        """)
        enxoval = cursor.fetchone()
        self.card_enxoval.config(text=f"{enxoval[0]}\n{formatar_moeda(enxoval[1] or 0)}")
        
        # Suprimentos repostos
        cursor.execute("""
            SELECT COUNT(*), SUM(valor_gasto)
            FROM reposicao_suprimentos
            WHERE date(data) >= date('now', '-7 days')
        """)
        suprimentos = cursor.fetchone()
        self.card_suprimentos.config(text=f"{suprimentos[0]}\n{formatar_moeda(suprimentos[1] or 0)}")
        
        # Total a receber
        total = (limpezas[1] or 0) + (enxoval[1] or 0) + (50.0 * cursor.execute("SELECT COUNT(DISTINCT imovel_id) FROM limpezas WHERE date(data) >= date('now', '-7 days')").fetchone()[0])
        self.card_receber.config(text=formatar_moeda(total))
        
        # Gerar gráficos
        self.gerar_grafico_limpezas()
        self.gerar_grafico_enxoval()
        
        conn.close()
    def gerar_grafico_limpezas(self):
        """Atualiza o gráfico de limpezas diretamente no canvas"""
        conn = sqlite3.connect("sistema.db")
        cursor = conn.cursor()

        cursor.execute("""
            SELECT date(data), COUNT(*)
            FROM limpezas
            WHERE date(data) >= date('now', '-30 days')
            GROUP BY date(data)
            ORDER BY date(data)
        """)

        datas = []
        quantidades = []

        for row in cursor.fetchall():
            datas.append(row[0][5:])  # Mostrar apenas dia/mês
            quantidades.append(row[1])

        # Limpar gráfico anterior
        fig = self.canvas_grafico1.figure
        fig.clear()
        ax = fig.add_subplot(111)

        # Gráfico moderno
        ax.plot(datas, quantidades, marker='o', linestyle='-', color=COR_DESTAQUE, linewidth=2)
        ax.fill_between(datas, quantidades, color=COR_DESTAQUE, alpha=0.2)
        ax.set_title('Limpezas nos Últimos 30 Dias', fontsize=10)
        ax.set_ylabel('Quantidade')
        ax.tick_params(axis='x', rotation=45, labelsize=8)
        fig.patch.set_facecolor(COR_CARD)
        ax.set_facecolor("#FFFFFF")
        ax.grid(True, linestyle="--", alpha=0.3)

        self.canvas_grafico1.draw()
        conn.close()

    
    def gerar_grafico_enxoval(self):
        """Atualiza o gráfico de enxoval diretamente no canvas"""
        conn = sqlite3.connect("sistema.db")
        cursor = conn.cursor()

        cursor.execute("""
            SELECT t.nome, SUM(c.quantidade)
            FROM consumo_enxoval c
            JOIN tipos_enxoval t ON c.item_id = t.id
            WHERE date(c.data) >= date('now', '-30 days')
            GROUP BY t.nome
            ORDER BY SUM(c.quantidade) DESC
            LIMIT 5
        """)

        itens = []
        quantidades = []

        for row in cursor.fetchall():
            itens.append(row[0])
            quantidades.append(row[1])

        fig = self.canvas_grafico2.figure
        fig.clear()
        ax = fig.add_subplot(111)

        # Gráfico de barras horizontal moderno
        bars = ax.barh(itens, quantidades, color=COR_SECUNDARIA)
        ax.set_title('Itens Mais Utilizados', fontsize=10)
        ax.set_xlabel('Quantidade')
        fig.patch.set_facecolor(COR_CARD)
        ax.set_facecolor("#FFFFFF")
        ax.grid(True, axis='x', linestyle="--", alpha=0.3)
        ax.bar_label(bars, fmt='%d', label_type='edge', padding=3)

        self.canvas_grafico2.draw()
        conn.close()

    def exibir_imagem_no_canvas(self, caminho_imagem, canvas):
        from PIL import Image, ImageTk

        # Acesse o widget Tkinter real dentro do FigureCanvasTkAgg
        tk_widget = canvas.get_tk_widget()

        # Obtenha dimensões reais do widget
        largura = tk_widget.winfo_width()
        altura = tk_widget.winfo_height()

        # Se ainda não estiver visível, espera o layout carregar
        if largura == 1 or altura == 1:
            # Agenda nova tentativa após 100ms
            self.frame_dashboard.after(100, lambda: self.exibir_imagem_no_canvas(caminho_imagem, canvas))
            return

        # Carrega e redimensiona a imagem
        img = Image.open(caminho_imagem)
        img = img.resize((largura, altura), Image.LANCZOS)
        img_tk = ImageTk.PhotoImage(img)

        # Armazena referência da imagem para evitar garbage collection
        canvas._img_ref = img_tk

        # Adiciona a imagem ao canvas
        canvas_widget = tk_widget
        label = Label(canvas_widget, image=img_tk, bg=COR_CARD)
        label.place(relx=0.5, rely=0.5, anchor='center')

    
    # =============================================
    # MÓDULO CLIENTES
    # =============================================
    
    def criar_clientes(self):
        """Cria a interface para gestão de clientes com layout moderno"""
        self.frame_clientes = Frame(self.frame_conteudo, bg=COR_FUNDO)
        
        # === Treeview para listar clientes ===
        frame_tree = Frame(self.frame_clientes, bg=COR_FUNDO)
        frame_tree.pack(fill=BOTH, expand=True, padx=10, pady=(10, 5))
        
        scrollbar = Scrollbar(frame_tree)
        scrollbar.pack(side=RIGHT, fill=Y)
        
        self.tree_clientes = ttk.Treeview(
            frame_tree, 
            columns=('nome', 'telefone', 'email'),
            yscrollcommand=scrollbar.set,
            show='headings',  # Remove a coluna #0 (desnecessária)
            height=10
        )
        self.tree_clientes.pack(fill=BOTH, expand=True)
        scrollbar.config(command=self.tree_clientes.yview)

        self.tree_clientes.heading('nome', text='Nome')
        self.tree_clientes.heading('telefone', text='Telefone')
        self.tree_clientes.heading('email', text='Email')
        
        self.tree_clientes.column('nome', width=150)
        self.tree_clientes.column('telefone', width=100)
        self.tree_clientes.column('email', width=200)

        # === Formulário de Cadastro ===
        frame_form = Frame(
            self.frame_clientes,
            bg=COR_CARD,
            bd=0,
            highlightthickness=1,
            highlightbackground="#e0e0e0"
        )
        frame_form.pack(fill=X, padx=10, pady=(5, 10))
        
        Label(
            frame_form,
            text="Cadastro de Clientes",
            bg=COR_CARD,
            fg=COR_TEXTO,
            font=self.fonte_titulo
        ).pack(pady=(10, 15), anchor="w", padx=10)
        
        # Campos do formulário em uma grade
        form_grid = Frame(frame_form, bg=COR_CARD)
        form_grid.pack(fill=X, padx=10)

        labels = ["Nome:", "Telefone:", "Email:", "Endereço:"]
        self.entry_cliente_nome = Entry(form_grid)
        self.entry_cliente_telefone = Entry(form_grid)
        self.entry_cliente_email = Entry(form_grid)
        self.entry_cliente_endereco = Entry(form_grid)
        entries = [
            self.entry_cliente_nome,
            self.entry_cliente_telefone,
            self.entry_cliente_email,
            self.entry_cliente_endereco
        ]
        
        for i, (label, entry) in enumerate(zip(labels, entries)):
            Label(
                form_grid, text=label,
                bg=COR_CARD,
                anchor="w"
            ).grid(row=i, column=0, sticky="w", padx=5, pady=5)
            
            entry.grid(row=i, column=1, sticky="ew", padx=5, pady=5)

        form_grid.columnconfigure(1, weight=1)

        # === Botões ===
        frame_botoes = Frame(frame_form, bg=COR_CARD)
        frame_botoes.pack(fill=X, padx=10, pady=(10, 15))
        
        estilo_botao = {
            "width": 12,
            "padx": 10,
            "pady": 5,
            "bd": 0,
            "font": ("Arial", 10, "bold"),
            "activebackground": "#cccccc"
        }

        Button(
            frame_botoes,
            text="Adicionar",
            command=self.adicionar_cliente,
            bg=COR_DESTAQUE,
            fg="white",
            **estilo_botao
        ).pack(side=LEFT, padx=(0, 10))

        Button(
            frame_botoes,
            text="Limpar",
            command=self.limpar_form_cliente,
            bg=COR_SECUNDARIA,
            fg="white",
            **estilo_botao
        ).pack(side=LEFT)

    
    def carregar_clientes(self):
        """Carrega os clientes no TreeView"""
        conn = sqlite3.connect("sistema.db")
        cursor = conn.cursor()
        cursor.execute("SELECT id, nome, telefone, email FROM clientes")
        
        # Limpar treeview
        for item in self.tree_clientes.get_children():
            self.tree_clientes.delete(item)
        
        # Adicionar novos itens
        for row in cursor.fetchall():
            self.tree_clientes.insert('', 'end', values=row)
        
        conn.close()
    
    def adicionar_cliente(self):
        """Adiciona um novo cliente ao banco de dados"""
        nome = self.entry_cliente_nome.get()
        telefone = self.entry_cliente_telefone.get()
        email = self.entry_cliente_email.get()
        endereco = self.entry_cliente_endereco.get()
        
        if nome:
            conn = sqlite3.connect("sistema.db")
            cursor = conn.cursor()
            cursor.execute("""
                INSERT INTO clientes (nome, telefone, email, endereco)
                VALUES (?, ?, ?, ?)
            """, (nome, telefone, email, endereco))
            conn.commit()
            conn.close()
            
            messagebox.showinfo("Sucesso", "Cliente cadastrado com sucesso!")
            self.carregar_clientes()
            self.limpar_form_cliente()
        else:
            messagebox.showerror("Erro", "Informe pelo menos o nome do cliente")
    
    def limpar_form_cliente(self):
        """Limpa o formulário de clientes"""
        self.entry_cliente_nome.delete(0, END)
        self.entry_cliente_telefone.delete(0, END)
        self.entry_cliente_email.delete(0, END)
        self.entry_cliente_endereco.delete(0, END)
    
    # =============================================
    # MÓDULO IMÓVEIS
    # =============================================
    def criar_imoveis(self):
        """Cria a interface para gestão de imóveis com visual moderno"""
        self.frame_imoveis = Frame(self.frame_conteudo, bg=COR_FUNDO)
        
        # === Treeview para listar imóveis ===
        frame_tree = Frame(self.frame_imoveis, bg=COR_FUNDO)
        frame_tree.pack(fill=BOTH, expand=True, padx=10, pady=(10, 5))

        scrollbar = Scrollbar(frame_tree)
        scrollbar.pack(side=RIGHT, fill=Y)

        self.tree_imoveis = ttk.Treeview(
            frame_tree,
            columns=('id', 'cliente', 'endereco', 'quartos', 'banheiros', 'plataforma'),
            show='headings',
            yscrollcommand=scrollbar.set,
            height=10
        )
        self.tree_imoveis.pack(fill=BOTH, expand=True)
        scrollbar.config(command=self.tree_imoveis.yview)

        # Esconder a coluna 'id' que usaremos apenas para referência
        self.tree_imoveis.column('id', width=0, stretch=NO)
        self.tree_imoveis.heading('id', text='ID')

        colunas = {
            'cliente': 'Cliente',
            'endereco': 'Endereço',
            'quartos': 'Quartos',
            'banheiros': 'Banheiros',
            'plataforma': 'Plataforma'
        }

        for col, nome in colunas.items():
            self.tree_imoveis.heading(col, text=nome)
            self.tree_imoveis.column(col, anchor='center', width=120)

        # Conecta o clique duplo para abrir edição do imóvel
        self.tree_imoveis.bind("<Double-1>", self.abrir_janela_edicao_imovel)

        # === Formulário de Cadastro de Imóveis ===
        frame_form = Frame(
            self.frame_imoveis,
            bg=COR_CARD,
            bd=0,
            highlightthickness=1,
            highlightbackground="#e0e0e0"
        )
        frame_form.pack(fill=X, padx=10, pady=(5, 10))
        
        Label(
            frame_form,
            text="Cadastro de Imóveis",
            bg=COR_CARD,
            fg=COR_TEXTO,
            font=self.fonte_titulo
        ).pack(pady=(10, 15), anchor="w", padx=10)

        form_grid = Frame(frame_form, bg=COR_CARD)
        form_grid.pack(fill=X, padx=10)

        # Cliente (Combobox)
        Label(form_grid, text="Cliente:", bg=COR_CARD, anchor="w", font=self.fonte_pequena).grid(
            row=0, column=0, sticky='w', padx=5, pady=5)
        self.combo_cliente_imovel = ttk.Combobox(form_grid, font=self.fonte_pequena)
        self.combo_cliente_imovel.grid(row=0, column=1, sticky='ew', padx=5, pady=5)

        # Endereço
        Label(form_grid, text="Endereço:", bg=COR_CARD, anchor="w", font=self.fonte_pequena).grid(
            row=1, column=0, sticky='w', padx=5, pady=5)
        self.entry_imovel_endereco = Entry(form_grid, font=self.fonte_pequena)
        self.entry_imovel_endereco.grid(row=1, column=1, sticky='ew', padx=5, pady=5)

        # Quartos
        Label(form_grid, text="Quartos:", bg=COR_CARD, anchor="w", font=self.fonte_pequena).grid(
            row=2, column=0, sticky='w', padx=5, pady=5)
        self.entry_imovel_quartos = Entry(form_grid, font=self.fonte_pequena)
        self.entry_imovel_quartos.grid(row=2, column=1, sticky='ew', padx=5, pady=5)

        # Banheiros
        Label(form_grid, text="Banheiros:", bg=COR_CARD, anchor="w", font=self.fonte_pequena).grid(
            row=3, column=0, sticky='w', padx=5, pady=5)
        self.entry_imovel_banheiros = Entry(form_grid, font=self.fonte_pequena)
        self.entry_imovel_banheiros.grid(row=3, column=1, sticky='ew', padx=5, pady=5)

        # Plataforma
        Label(form_grid, text="Plataforma:", bg=COR_CARD, anchor="w", font=self.fonte_pequena).grid(
            row=4, column=0, sticky='w', padx=5, pady=5)
        self.entry_imovel_plataforma = Entry(form_grid, font=self.fonte_pequena)
        self.entry_imovel_plataforma.grid(row=4, column=1, sticky='ew', padx=5, pady=5)

        form_grid.columnconfigure(1, weight=1)

        # === Botões ===
        frame_botoes = Frame(frame_form, bg=COR_CARD)
        frame_botoes.pack(fill=X, padx=10, pady=(10, 15))
        
        estilo_botao = {
            "width": 12,
            "padx": 10,
            "pady": 5,
            "bd": 0,
            "font": ("Arial", 10, "bold"),
            "activebackground": "#cccccc"
        }

        Button(
            frame_botoes,
            text="Adicionar",
            command=self.adicionar_imovel,
            bg=COR_DESTAQUE,
            fg="white",
            **estilo_botao
        ).pack(side=LEFT, padx=(0, 10))

        Button(
            frame_botoes,
            text="Limpar",
            command=self.limpar_form_imovel,
            bg=COR_SECUNDARIA,
            fg="white",
            **estilo_botao
        ).pack(side=LEFT)
        
        # Carregar os imóveis do banco de dados
        self.carregar_imoveis()

    def adicionar_imovel(self):
        """Adiciona um imóvel ao banco de dados e atualiza a Treeview"""
        cliente = self.combo_cliente_imovel.get().strip()
        endereco = self.entry_imovel_endereco.get().strip()
        quartos = self.entry_imovel_quartos.get().strip()
        banheiros = self.entry_imovel_banheiros.get().strip()
        plataforma = self.entry_imovel_plataforma.get().strip()

        # Validação simples
        if not (cliente and endereco and quartos and banheiros and plataforma):
            messagebox.showwarning("Campos obrigatórios", "Por favor, preencha todos os campos.")
            return

        try:
            # Extrai o ID do cliente do combobox (formato "ID - Nome")
            cliente_id = int(cliente.split(' - ')[0])
            quartos = int(quartos)
            banheiros = int(banheiros)
        except (ValueError, IndexError):
            messagebox.showerror("Erro de entrada", "Verifique os dados inseridos.")
            return

        # Inserir no banco de dados
        conn = sqlite3.connect("sistema.db")
        cursor = conn.cursor()
        
        try:
            cursor.execute("""
                INSERT INTO imoveis 
                (cliente_id, endereco, quartos, banheiros, plataforma)
                VALUES (?, ?, ?, ?, ?)
            """, (cliente_id, endereco, quartos, banheiros, plataforma))
            conn.commit()
        except sqlite3.Error as e:
            messagebox.showerror("Erro no banco de dados", f"Não foi possível salvar o imóvel:\n{str(e)}")
        finally:
            conn.close()

        # Recarregar a lista de imóveis
        self.carregar_imoveis()
        
        # Limpar formulário
        self.limpar_form_imovel()
        
        messagebox.showinfo("Sucesso", "Imóvel cadastrado com sucesso!")

    def carregar_imoveis(self):
        """Carrega os imóveis no TreeView e atualiza o combobox de clientes"""
        conn = sqlite3.connect("sistema.db")
        cursor = conn.cursor()
        
        # Verifica se a coluna plataforma existe e a adiciona se necessário
        cursor.execute("PRAGMA table_info(imoveis)")
        colunas = [col[1] for col in cursor.fetchall()]
        if 'plataforma' not in colunas:
            cursor.execute("ALTER TABLE imoveis ADD COLUMN plataforma TEXT")
            conn.commit()
        
        # Restante do código permanece o mesmo...
        cursor.execute("SELECT id, nome FROM clientes")
        clientes = cursor.fetchall()
        self.combo_cliente_imovel['values'] = [f"{c[0]} - {c[1]}" for c in clientes]
        
        if clientes:
            self.combo_cliente_imovel.current(0)
        
        cursor.execute("""
            SELECT i.id, c.nome, i.endereco, i.quartos, i.banheiros, 
                COALESCE(i.plataforma, '') as plataforma
            FROM imoveis i
            JOIN clientes c ON i.cliente_id = c.id
        """)
        
        for item in self.tree_imoveis.get_children():
            self.tree_imoveis.delete(item)
        
        for row in cursor.fetchall():
            self.tree_imoveis.insert('', 'end', values=row)
        
        conn.close()

    def limpar_form_imovel(self):
            """Limpa o formulário de imóveis"""
            self.combo_cliente_imovel.set('')  # Limpa o combobox
            self.entry_imovel_endereco.delete(0, END)
            self.entry_imovel_quartos.delete(0, END)
            self.entry_imovel_banheiros.delete(0, END)
            self.entry_imovel_plataforma.delete(0, END)

    def abrir_janela_edicao_imovel(self, event):
        """Abre uma janela para editar o imóvel selecionado"""
        # Pega o item selecionado na Treeview
        item_selecionado = self.tree_imoveis.selection()
        if not item_selecionado:
            return  # Nada selecionado, sai da função
        
        item = item_selecionado[0]
        dados = self.tree_imoveis.item(item, "values")
        
        # Cria a janela de edição
        janela_edicao = Toplevel()
        janela_edicao.title("Editar Imóvel")
        janela_edicao.resizable(False, False)
        
        # Frame principal
        frame_principal = Frame(janela_edicao, padx=20, pady=20)
        frame_principal.pack()
        
        # Variáveis para os campos
        var_endereco = StringVar(value=dados[2])
        var_quartos = StringVar(value=dados[3])
        var_banheiros = StringVar(value=dados[4])
        var_plataforma = StringVar(value=dados[5])
        
        # Campos do formulário
        Label(frame_principal, text="Endereço:").grid(row=0, column=0, sticky='w', pady=5)
        Entry(frame_principal, textvariable=var_endereco, width=40).grid(row=0, column=1, pady=5)
        
        Label(frame_principal, text="Quartos:").grid(row=1, column=0, sticky='w', pady=5)
        Entry(frame_principal, textvariable=var_quartos, width=10).grid(row=1, column=1, sticky='w', pady=5)
        
        Label(frame_principal, text="Banheiros:").grid(row=2, column=0, sticky='w', pady=5)
        Entry(frame_principal, textvariable=var_banheiros, width=10).grid(row=2, column=1, sticky='w', pady=5)
        
        Label(frame_principal, text="Plataforma:").grid(row=3, column=0, sticky='w', pady=5)
        Entry(frame_principal, textvariable=var_plataforma, width=20).grid(row=3, column=1, sticky='w', pady=5)
        
        # Frame para botões
        frame_botoes = Frame(frame_principal)
        frame_botoes.grid(row=4, column=0, columnspan=2, pady=10)
    
   
    # Função para salvar as alterações
    def salvar_edicao():
        # Validação dos campos
        if not (var_endereco.get() and var_quartos.get() and var_banheiros.get() and var_plataforma.get()):
            messagebox.showwarning("Campos obrigatórios", "Preencha todos os campos!")
            return
        
        try:
            quartos = int(var_quartos.get())
            banheiros = int(var_banheiros.get())
        except ValueError:
            messagebox.showerror("Erro", "Quartos e banheiros devem ser números!")
            return
        
        # Atualizar no banco de dados
        conn = sqlite3.connect("sistema.db")
        cursor = conn.cursor()
        try:
            cursor.execute("""
                UPDATE imoveis 
                SET endereco=?, quartos=?, banheiros=?, plataforma=?
                WHERE id=?
            """, (var_endereco.get(), quartos, banheiros, var_plataforma.get(), dados[0]))
            conn.commit()
            messagebox.showinfo("Sucesso", "Imóvel atualizado com sucesso!")
            
            # Atualizar a Treeview
            self.carregar_imoveis()
            janela_edicao.destroy()
        except sqlite3.Error as e:
            messagebox.showerror("Erro", f"Não foi possível atualizar:\n{str(e)}")
        finally:
            conn.close()
    
        # Botões
        Button(frame_botoes, text="Salvar", command=salvar_edicao, width=10).pack(side=LEFT, padx=5)
        Button(frame_botoes, text="Cancelar", command=janela_edicao.destroy, width=10).pack(side=LEFT, padx=5)
    #========================================
    # MÓDULO LIMPEZA
    # =============================================
    
    def criar_limpeza(self):
        """Cria a interface para gestão de limpezas"""
        self.frame_limpeza = Frame(self.frame_conteudo, bg=COR_FUNDO)
        
        # Treeview para listar limpezas
        frame_tree = Frame(self.frame_limpeza, bg=COR_FUNDO)
        frame_tree.pack(fill=BOTH, expand=True, padx=10, pady=(10, 5))
        
        scrollbar = Scrollbar(frame_tree)
        scrollbar.pack(side=RIGHT, fill=Y)
        
        self.tree_limpezas = ttk.Treeview(frame_tree, 
                                        columns=('id', 'imovel', 'data', 'horas', 'valor_hora', 'valor_total'),
                                        yscrollcommand=scrollbar.set)
        self.tree_limpezas.pack(fill=BOTH, expand=True)
        scrollbar.config(command=self.tree_limpezas.yview)
        
        self.tree_limpezas.heading('#0', text='ID')
        self.tree_limpezas.heading('#1', text='Imóvel')
        self.tree_limpezas.heading('#2', text='Data')
        self.tree_limpezas.heading('#3', text='Horas')
        self.tree_limpezas.heading('#4', text='Valor/Hora')
        self.tree_limpezas.heading('#5', text='Valor Total')
        
        # Formulário para adicionar limpezas
        frame_form = Frame(self.frame_limpeza, bg=COR_CARD, bd=0, 
                        highlightthickness=1, highlightbackground="#e0e0e0",
                        padx=10, pady=10)
        frame_form.pack(fill=X, padx=10, pady=(5, 10))
        
        # Título do formulário
        Label(frame_form, text="Registro de Limpeza", bg=COR_CARD, 
            fg=COR_TEXTO, font=self.fonte_titulo).grid(row=0, column=0, columnspan=3, 
                                                        pady=(0, 15), sticky="w")
        
        # Grid para organização dos campos
        rows = [
            ("Imóvel:", ttk.Combobox(frame_form)),
            ("Data:", DateEntry(frame_form, date_pattern='dd/mm/yyyy')),
            ("Hora Início:", Entry(frame_form)),
            ("Hora Fim:", Entry(frame_form)),
            ("Valor por Hora (R$):", Entry(frame_form)),
            ("Observações:", Text(frame_form, height=3))
        ]
        
        # Configurar widgets e referências
        self.combo_imovel_limpeza = rows[0][1]
        self.entry_limpeza_data = rows[1][1]
        self.entry_limpeza_hora_inicio = rows[2][1]
        self.entry_limpeza_hora_fim = rows[3][1]
        self.entry_limpeza_valor_hora = rows[4][1]
        self.entry_limpeza_observacoes = rows[5][1]
        
        # Configurar valor padrão para hora
        self.entry_limpeza_valor_hora.insert(0, "30.00")
        
        # Posicionar widgets usando grid
        for i, (label_text, widget) in enumerate(rows, start=1):
            # Label
            Label(frame_form, text=label_text, bg=COR_CARD, anchor="e", width=15).grid(
                row=i, column=0, padx=(0, 5), pady=3, sticky="e")
            
            # Widget
            if isinstance(widget, Text):
                widget.grid(row=i, column=1, columnspan=2, sticky="ew", pady=3)
                # Adicionar scrollbar para o Text
                scroll_text = Scrollbar(frame_form, command=widget.yview)
                scroll_text.grid(row=i, column=3, sticky="ns")
                widget.config(yscrollcommand=scroll_text.set)
            else:
                widget.grid(row=i, column=1, columnspan=2, sticky="ew", pady=3)
        
        # Configurar pesos das colunas
        frame_form.columnconfigure(0, weight=0)
        frame_form.columnconfigure(1, weight=1)
        frame_form.columnconfigure(2, weight=1)
        
        # Frame para botões (centralizado)
        frame_botoes = Frame(frame_form, bg=COR_CARD)
        frame_botoes.grid(row=len(rows)+1, column=0, columnspan=3, pady=(10, 0))
        
        # Botões com espaçamento uniforme
        Button(frame_botoes, text="Adicionar", command=self.adicionar_limpeza,
            bg=COR_DESTAQUE, fg="white", width=12).pack(side=LEFT, padx=5, ipady=3)
        Button(frame_botoes, text="Calcular", command=self.calcular_limpeza,
            bg=COR_SECUNDARIA, fg="white", width=12).pack(side=LEFT, padx=5, ipady=3)
        Button(frame_botoes, text="Limpar", command=self.limpar_form_limpeza,
            bg=COR_ALERTA, fg="white", width=12).pack(side=LEFT, padx=5, ipady=3)
        
        # Configurar padding para todos os widgets filhos
        for child in frame_form.winfo_children():
            child.grid_configure(padx=5, pady=2)
            
    def carregar_limpezas(self):
        """Carrega as limpezas no TreeView e atualiza o combobox de imóveis"""
        conn = sqlite3.connect("sistema.db")
        cursor = conn.cursor()
        
        # Carregar imóveis no combobox
        cursor.execute("SELECT id, endereco FROM imoveis")
        imoveis = cursor.fetchall()
        self.combo_imovel_limpeza['values'] = [f"{i[0]} - {i[1]}" for i in imoveis]
        
        if imoveis:
            self.combo_imovel_limpeza.current(0)
        
        # Carregar limpezas no treeview
        cursor.execute("""
        CREATE TABLE IF NOT EXISTS imoveis (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        cliente_id INTEGER,
        endereco TEXT NOT NULL,
        quartos INTEGER,
        banheiros INTEGER,
        FOREIGN KEY (cliente_id) REFERENCES clientes (id)
    );
    """)
        
        # Limpar treeview
        for item in self.tree_limpezas.get_children():
            self.tree_limpezas.delete(item)
        
        # Adicionar novos itens formatados
        for row in cursor.fetchall():
            self.tree_limpezas.insert('', 'end', values=(
                row[0], row[1], row[2], 
                f"{row[3]:.2f}h", 
                formatar_moeda(row[4]), 
                formatar_moeda(row[5])
            ))
        
        conn.close()
    
    def calcular_limpeza(self):
        """Calcula o valor da limpeza com base nas horas trabalhadas"""
        hora_inicio = self.entry_limpeza_hora_inicio.get()
        hora_fim = self.entry_limpeza_hora_fim.get()
        valor_hora = self.entry_limpeza_valor_hora.get()
        
        try:
            horas = calcular_horas(hora_inicio, hora_fim)
            valor = horas * float(valor_hora)
            messagebox.showinfo("Cálculo", 
                              f"Horas trabalhadas: {horas:.2f}h\n"
                              f"Valor total: {formatar_moeda(valor)}")
        except ValueError:
            messagebox.showerror("Erro", "Informe horários no formato HH:MM e valor por hora numérico")
    
    def adicionar_limpeza(self):
        """Adiciona uma nova limpeza ao banco de dados"""
        imovel_id = self.combo_imovel_limpeza.get().split(" - ")[0]
        data = self.entry_limpeza_data.get_date()
        hora_inicio = self.entry_limpeza_hora_inicio.get()
        hora_fim = self.entry_limpeza_hora_fim.get()
        valor_hora = self.entry_limpeza_valor_hora.get()
        observacoes = self.entry_limpeza_observacoes.get("1.0", END).strip()
        
        if imovel_id and data and hora_inicio and hora_fim and valor_hora:
            try:
                horas = calcular_horas(hora_inicio, hora_fim)
                valor_total = horas * float(valor_hora)
                
                conn = sqlite3.connect("sistema.db")
                cursor = conn.cursor()
                cursor.execute("""
                    INSERT INTO limpezas 
                    (imovel_id, data, hora_inicio, hora_fim, horas_trabalhadas, valor_hora, valor_total, observacoes)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                """, (imovel_id, data, hora_inicio, hora_fim, horas, float(valor_hora), valor_total, observacoes))
                conn.commit()
                conn.close()
                
                messagebox.showinfo("Sucesso", "Limpeza registrada com sucesso!")
                self.carregar_limpezas()
                self.limpar_form_limpeza()
            except ValueError:
                messagebox.showerror("Erro", "Valor por hora deve ser numérico")
        else:
            messagebox.showerror("Erro", "Preencha todos os campos obrigatórios")
    
    def limpar_form_limpeza(self):
        """Limpa o formulário de limpeza"""
        self.entry_limpeza_hora_inicio.delete(0, END)
        self.entry_limpeza_hora_fim.delete(0, END)
        self.entry_limpeza_observacoes.delete("1.0", END)
    
    # =============================================
    # MÓDULO ENXOVAL
    # =============================================
    
    def criar_enxoval(self):
        """Cria a interface para gestão de enxoval com layout profissional"""
        self.frame_enxoval = Frame(self.frame_conteudo, bg=COR_FUNDO)
        
        # Container principal com grid layout
        container = Frame(self.frame_enxoval, bg=COR_FUNDO)
        container.pack(fill=BOTH, expand=True, padx=20, pady=20)
        
        # Seção da tabela (60% da largura)
        frame_tabela = Frame(container, bg=COR_FUNDO)
        frame_tabela.grid(row=0, column=0, sticky="nsew", padx=(0, 15))
        
        # Treeview para listar consumo de enxoval
        Label(frame_tabela, text="Histórico de Consumo", bg=COR_FUNDO, 
            fg=COR_TEXTO, font=self.fonte_subtitulo).pack(anchor="w", pady=(0, 10))
        
        frame_tree = Frame(frame_tabela, bg=COR_FUNDO)
        frame_tree.pack(fill=BOTH, expand=True)
        
        scroll_y = Scrollbar(frame_tree)
        scroll_y.pack(side=RIGHT, fill=Y)
        
        scroll_x = Scrollbar(frame_tree, orient='horizontal')
        scroll_x.pack(side=BOTTOM, fill=X)
        
        self.tree_enxoval = ttk.Treeview(
            frame_tree,
            columns=('id', 'imovel', 'item', 'quantidade', 'data', 'valor_unitario', 'valor_total'),
            yscrollcommand=scroll_y.set,
            xscrollcommand=scroll_x.set,
            selectmode='browse',
            height=15
        )
        self.tree_enxoval.pack(fill=BOTH, expand=True)
        
        scroll_y.config(command=self.tree_enxoval.yview)
        scroll_x.config(command=self.tree_enxoval.xview)
        
        # Configuração das colunas
        colunas = [
            ('ID', 50, 'center'),
            ('Imóvel', 150, 'w'),
            ('Item', 120, 'w'),
            ('Quantidade', 80, 'center'),
            ('Data', 100, 'center'),
            ('Valor Unitário', 100, 'e'),
            ('Valor Total', 120, 'e')
        ]
        
        for idx, (text, width, anchor) in enumerate(colunas):
            self.tree_enxoval.heading(f'#{idx}', text=text)
            self.tree_enxoval.column(f'#{idx}', width=width, anchor=anchor)
        
        # Remover a coluna #0 extra
        self.tree_enxoval.column('#0', width=0, stretch=NO)
        
        # Seção do formulário (40% da largura)
        frame_formulario = Frame(container, bg=COR_FUNDO)
        frame_formulario.grid(row=0, column=1, sticky="nsew")
        
        # Card do formulário
        form_card = Frame(frame_formulario, bg=COR_CARD, padx=15, pady=15,
                        highlightthickness=1, highlightbackground="#e0e0e0")
        form_card.pack(fill=BOTH, expand=True)
        
        Label(form_card, text="Registrar Consumo", bg=COR_CARD, 
            fg=COR_TEXTO, font=self.fonte_subtitulo).grid(row=0, column=0, columnspan=2, pady=(0, 15), sticky="w")
        
        # Campos do formulário com grid layout
        campos = [
            ('Imóvel:', self.combo_imovel_enxoval, ttk.Combobox(form_card)),
            ('Item:', self.combo_item_enxoval, ttk.Combobox(form_card)),
            ('Quantidade:', self.entry_enxoval_quantidade, Entry(form_card)),
            ('Data:', self.entry_enxoval_data, DateEntry(form_card, date_pattern='dd/mm/yyyy'))
        ]
        
        for row, (label_text, var, widget) in enumerate(campos, start=1):
            # Atribui o widget à variável de instância
            setattr(self, var if isinstance(var, str) else var.widget, widget)
            
            Label(form_card, text=label_text, bg=COR_CARD).grid(
                row=row, column=0, padx=5, pady=5, sticky="w")
            
            widget.grid(row=row, column=1, padx=5, pady=5, sticky="ew")
            form_card.grid_columnconfigure(1, weight=1)  # Expande a coluna dos campos
        
        # Frame de botões
        frame_botoes = Frame(form_card, bg=COR_CARD)
        frame_botoes.grid(row=len(campos)+1, column=0, columnspan=2, pady=(15, 0), sticky="ew")
        
        Button(frame_botoes, text="Adicionar", command=self.adicionar_consumo_enxoval,
            bg=COR_DESTAQUE, fg="white", width=12).pack(side=LEFT, padx=5)
        Button(frame_botoes, text="Limpar", command=self.limpar_form_enxoval,
            bg=COR_SECUNDARIA, fg="white", width=12).pack(side=LEFT, padx=5)
        
        # Configuração do grid do container principal
        container.grid_columnconfigure(0, weight=6)  # 60% para tabela
        container.grid_columnconfigure(1, weight=4)  # 40% para formulário
        container.grid_rowconfigure(0, weight=1)
        
        # Carregar dados iniciais
        self.carregar_itens_enxoval()

# Os métodos adicionar_consumo_enxoval, carregar_itens_enxoval e limpar_form_enxoval permanecem os mesmos

    def adicionar_consumo_enxoval(self):
            """Adiciona um novo consumo de enxoval ao banco de dados"""
            try:
                imovel = self.combo_imovel_enxoval.get()
                item = self.combo_item_enxoval.get()
                quantidade = self.entry_enxoval_quantidade.get()
                data = self.entry_enxoval_data.get_date()
                
                if not all([imovel, item, quantidade, data]):
                    messagebox.showerror("Erro", "Preencha todos os campos obrigatórios!")
                    return
                    
                imovel_id = imovel.split(" - ")[0]
                item_id = item.split(" - ")[0]
                quantidade = int(quantidade)
                
                conn = sqlite3.connect("sistema.db")
                cursor = conn.cursor()
                
                # Verificar se o item existe
                cursor.execute("SELECT preco_unitario FROM tipos_enxoval WHERE id=?", (item_id,))
                resultado = cursor.fetchone()
                
                if not resultado:
                    messagebox.showerror("Erro", "Item selecionado não encontrado!")
                    return
                    
                preco_unitario = resultado[0]
                valor_total = quantidade * preco_unitario
                
                cursor.execute("""
                    INSERT INTO consumo_enxoval 
                    (imovel_id, item_id, quantidade, data)
                    VALUES (?, ?, ?, ?)
                """, (imovel_id, item_id, quantidade, data))
                
                conn.commit()
                conn.close()
                
                messagebox.showinfo("Sucesso", "Consumo de enxoval registrado com sucesso!")
                self.carregar_itens_enxoval()
                self.limpar_form_enxoval()
                
            except ValueError:
                messagebox.showerror("Erro", "Quantidade deve ser um número inteiro válido!")
            except Exception as e:
                messagebox.showerror("Erro", f"Ocorreu um erro ao registrar o consumo:\n{str(e)}")

    def carregar_itens_enxoval(self):
        """Carrega os itens de enxoval no TreeView e atualiza os comboboxes"""
        try:
            conn = sqlite3.connect("sistema.db")
            cursor = conn.cursor()
            
            # Carregar imóveis no combobox
            cursor.execute("SELECT id, endereco FROM imoveis")
            imoveis = cursor.fetchall()
            self.combo_imovel_enxoval['values'] = [f"{i[0]} - {i[1]}" for i in imoveis]
            
            if imoveis:
                self.combo_imovel_enxoval.current(0)
            
            # Carregar itens de enxoval no combobox
            cursor.execute("SELECT id, nome, preco_unitario FROM tipos_enxoval")
            itens = cursor.fetchall()
            self.combo_item_enxoval['values'] = [f"{i[0]} - {i[1]} (R$ {i[2]:.2f})" for i in itens]
            
            if itens:
                self.combo_item_enxoval.current(0)
            
            # Carregar consumo de enxoval no treeview
            cursor.execute("""
                SELECT c.id, i.endereco, t.nome, c.quantidade, c.data, t.preco_unitario, 
                    (c.quantidade * t.preco_unitario) as valor_total
                FROM consumo_enxoval c
                JOIN imoveis i ON c.imovel_id = i.id
                JOIN tipos_enxoval t ON c.item_id = t.id
                ORDER BY c.data DESC
            """)
            
            # Limpar treeview
            for item in self.tree_enxoval.get_children():
                self.tree_enxoval.delete(item)
            
            # Adicionar novos itens formatados
            for row in cursor.fetchall():
                self.tree_enxoval.insert('', 'end', values=(
                    row[0], row[1], row[2], row[3], row[4],
                    formatar_moeda(row[5]), 
                    formatar_moeda(row[6])
                ))
            
        except Exception as e:
            messagebox.showerror("Erro", f"Erro ao carregar itens de enxoval:\n{str(e)}")
        finally:
            conn.close()


    def limpar_form_enxoval(self):
        """Limpa o formulário de enxoval"""
        self.entry_enxoval_quantidade.delete(0, END)
        # Mantém os comboboxes com os valores atuais
        if self.combo_imovel_enxoval['values']:
            self.combo_imovel_enxoval.current(0)
        if self.combo_item_enxoval['values']:
            self.combo_item_enxoval.current(0)
        # Define a data atual como padrão
        self.entry_enxoval_data.set_date(datetime.now().date())
    
    def criar_config_itens(self):
        """Cria a interface para configuração de itens de enxoval e suprimentos"""
        self.frame_config_itens = Frame(self.frame_conteudo, bg=COR_FUNDO)
        # Abas para enxoval e suprimentos
        abas = ttk.Notebook(self.frame_config_itens)
        frame_enxoval = Frame(abas, bg=COR_FUNDO)
        frame_suprimentos = Frame(abas, bg=COR_FUNDO)
        abas.add(frame_enxoval, text="Itens de Enxoval")
        abas.add(frame_suprimentos, text="Itens de Suprimentos")
        abas.pack(fill=BOTH, expand=True)
        # Constrói as abas usando seu método auxiliar
        self._construir_aba_itens(frame_enxoval, "enxoval")
        self._construir_aba_itens(frame_suprimentos, "suprimentos")
        # Carrega os itens
        self.carregar_itens_config()

    def _construir_aba_itens(self, parent_frame, tipo_item):
        """Método auxiliar para construir a interface de cada tipo de item"""
        # Treeview para listar itens
        frame_tree = Frame(parent_frame, bg=COR_FUNDO)
        frame_tree.pack(fill=BOTH, expand=True, padx=10, pady=10)
        
        scrollbar = Scrollbar(frame_tree)
        scrollbar.pack(side=RIGHT, fill=Y)
        
        tree = ttk.Treeview(frame_tree, columns=('id', 'nome', 'preco', 'unidade'), 
                            yscrollcommand=scrollbar.set, selectmode='browse')
        tree.pack(fill=BOTH, expand=True)
        scrollbar.config(command=tree.yview)
        
        tree.heading('#0', text='ID')
        tree.heading('#1', text='Nome')
        tree.heading('#2', text='Preço Unitário')
        tree.heading('#3', text='Unidade')
        
        # Armazenar referência para acesso posterior
        if tipo_item == "enxoval":
            self.tree_enxoval_config = tree
        else:
            self.tree_suprimentos_config = tree
        
        # Configurar bind para seleção
        tree.bind('<<TreeviewSelect>>', lambda e: self._preencher_form_item(tipo_item))
        
        # Formulário para adicionar/editar itens
        frame_form = Frame(parent_frame, bg=COR_CARD, bd=0, 
                        highlightthickness=1, highlightbackground="#e0e0e0")
        frame_form.pack(fill=X, padx=10, pady=10)
        
        Label(frame_form, text=f"Configurar Itens de {tipo_item.capitalize()}", bg=COR_CARD, 
            fg=COR_TEXTO, font=self.fonte_titulo).pack(pady=(10, 5), anchor="w", padx=10)
        
        # Campos do formulário
        campos = [
            ("Nome:", Entry(frame_form)),
            ("Preço Unitário (R$):", Entry(frame_form)),
            ("Unidade de Medida:", ttk.Combobox(frame_form, values=['unidade', 'par', 'jogo', 'conjunto', 'metro']))
        ]
        
        # Armazenar referências aos widgets
        if tipo_item == "enxoval":
            self.entry_enxoval_nome = campos[0][1]
            self.entry_enxoval_preco = campos[1][1]
            self.combo_enxoval_unidade = campos[2][1]
        else:
            self.entry_suprimento_nome = campos[0][1]
            self.entry_suprimento_preco = campos[1][1]
            self.combo_suprimento_unidade = campos[2][1]
        
        for texto, widget in campos:
            frame = Frame(frame_form, bg=COR_CARD)
            frame.pack(fill=X, padx=10, pady=5)
            Label(frame, text=texto, bg=COR_CARD).pack(side=LEFT, padx=5)
            widget.pack(side=LEFT, expand=True, fill=X)
        
        # Botões
        frame_botoes = Frame(frame_form, bg=COR_CARD)
        frame_botoes.pack(fill=X, padx=10, pady=10)
        
        Button(frame_botoes, text="Salvar", 
            command=lambda: self._salvar_item(tipo_item),
            bg=COR_DESTAQUE, fg="white").pack(side=LEFT, padx=5)
        Button(frame_botoes, text="Remover", 
            command=lambda: self._remover_item(tipo_item),
            bg=COR_ALERTA, fg="white").pack(side=LEFT, padx=5)
        Button(frame_botoes, text="Limpar", 
            command=lambda: self._limpar_form_item(tipo_item),
            bg=COR_SECUNDARIA, fg="white").pack(side=LEFT, padx=5)

    def _preencher_form_item(self, tipo_item):
        """Preenche o formulário com o item selecionado"""
        if tipo_item == "enxoval":
            tree = self.tree_enxoval_config
            nome_entry = self.entry_enxoval_nome
            preco_entry = self.entry_enxoval_preco
            unidade_combo = self.combo_enxoval_unidade
        else:
            tree = self.tree_suprimentos_config
            nome_entry = self.entry_suprimento_nome
            preco_entry = self.entry_suprimento_preco
            unidade_combo = self.combo_suprimento_unidade
        
        selected = tree.selection()
        if selected:
            item = tree.item(selected[0])
            values = item['values']
            
            nome_entry.delete(0, END)
            nome_entry.insert(0, values[1])
            
            preco_entry.delete(0, END)
            preco_entry.insert(0, values[2].replace('R$ ', ''))
            
            unidade_combo.set(values[3])
            
            # Armazenar ID do item selecionado
            self.item_selecionado_id = values[0]
            self.item_selecionado_tipo = tipo_item

    def _salvar_item(self, tipo_item):
        """Salva ou atualiza um item"""
        if tipo_item == "enxoval":
            nome = self.entry_enxoval_nome.get()
            preco = self.entry_enxoval_preco.get()
            unidade = self.combo_enxoval_unidade.get()
            tabela = "tipos_enxoval"
        else:
            nome = self.entry_suprimento_nome.get()
            preco = self.entry_suprimento_preco.get()
            unidade = self.combo_suprimento_unidade.get()
            tabela = "suprimentos"
        
        if not nome:
            messagebox.showerror("Erro", "O nome é obrigatório!")
            return
        
        try:
            preco = float(preco) if preco else 0.0
        except ValueError:
            messagebox.showerror("Erro", "Preço deve ser um número válido!")
            return
        
        conn = sqlite3.connect("sistema.db")
        cursor = conn.cursor()
        
        if hasattr(self, 'item_selecionado_id') and self.item_selecionado_id and self.item_selecionado_tipo == tipo_item:
            # Atualizar item existente
            cursor.execute(f"""
                UPDATE {tabela} 
                SET nome=?, preco_unitario=?, unidade_medida=?
                WHERE id=?
            """, (nome, preco, unidade, self.item_selecionado_id))
        else:
            # Inserir novo item
            cursor.execute(f"""
                INSERT INTO {tabela} (nome, preco_unitario, unidade_medida)
                VALUES (?, ?, ?)
            """, (nome, preco, unidade))
        
        conn.commit()
        conn.close()
        
        messagebox.showinfo("Sucesso", "Item salvo com sucesso!")
        self.carregar_itens_config()
        self._limpar_form_item(tipo_item)

    def _remover_item(self, tipo_item):
        """Remove o item selecionado"""
        if not hasattr(self, 'item_selecionado_id') or not self.item_selecionado_id or self.item_selecionado_tipo != tipo_item:
            messagebox.showerror("Erro", "Nenhum item selecionado!")
            return
        
        resposta = messagebox.askyesno("Confirmar", "Tem certeza que deseja remover este item?")
        if resposta:
            tabela = "tipos_enxoval" if tipo_item == "enxoval" else "suprimentos"
            
            conn = sqlite3.connect("sistema.db")
            cursor = conn.cursor()
            
            try:
                # Verificar se o item está em uso
                if tipo_item == "enxoval":
                    cursor.execute("SELECT COUNT(*) FROM consumo_enxoval WHERE item_id=?", (self.item_selecionado_id,))
                else:
                    cursor.execute("SELECT COUNT(*) FROM consumo_suprimentos WHERE item_id=?", (self.item_selecionado_id,))
                
                if cursor.fetchone()[0] > 0:
                    messagebox.showerror("Erro", "Este item está em uso e não pode ser removido!")
                    return
                
                # Remover o item
                cursor.execute(f"DELETE FROM {tabela} WHERE id=?", (self.item_selecionado_id,))
                conn.commit()
                messagebox.showinfo("Sucesso", "Item removido!")
                
                self.carregar_itens_config()
                self._limpar_form_item(tipo_item)
            finally:
                conn.close()

    def _limpar_form_item(self, tipo_item):
            """Limpa o formulário de itens"""
            if tipo_item == "enxoval":
                self.entry_enxoval_nome.delete(0, END)
                self.entry_enxoval_preco.delete(0, END)
                self.combo_enxoval_unidade.set('')
            else:
                self.entry_suprimento_nome.delete(0, END)
                self.entry_suprimento_preco.delete(0, END)
                self.combo_suprimento_unidade.set('')
            
            if hasattr(self, 'item_selecionado_id'):
                del self.item_selecionado_id
            if hasattr(self, 'item_selecionado_tipo'):
                del self.item_selecionado_tipo

    def carregar_itens_config(self):
        """Carrega todos os itens para configuração"""
        conn = sqlite3.connect("sistema.db")
        cursor = conn.cursor()
        
        try:
            # Verificar se a coluna unidade_medida existe, se não, adicionar
            cursor.execute("PRAGMA table_info(tipos_enxoval)")
            colunas = [col[1] for col in cursor.fetchall()]
            
            if 'unidade_medida' not in colunas:
                cursor.execute("ALTER TABLE tipos_enxoval ADD COLUMN unidade_medida TEXT DEFAULT 'unidade'")
                conn.commit()
            
            # Carregar itens de enxoval
            cursor.execute("""
                SELECT id, nome, preco_unitario, 
                    COALESCE(unidade_medida, 'unidade') as unidade_medida 
                FROM tipos_enxoval 
                ORDER BY nome
            """)
            self._preencher_treeview(self.tree_enxoval_config, cursor.fetchall())
            
            # Carregar suprimentos (se aplicável)
            if hasattr(self, 'tree_suprimentos_config'):
                cursor.execute("""
                    SELECT id, nome, preco_unitario, 
                        COALESCE(unidade_medida, 'unidade') as unidade_medida 
                    FROM suprimentos 
                    ORDER BY nome
                """)
                self._preencher_treeview(self.tree_suprimentos_config, cursor.fetchall())
                
        except Exception as e:
            messagebox.showerror("Erro", f"Erro ao carregar itens:\n{str(e)}")
        finally:
            conn.close()

    def _preencher_treeview(self, tree, dados):
            """Preenche um treeview com os dados fornecidos"""
            # Limpar treeview
            for item in tree.get_children():
                tree.delete(item)
            
            # Adicionar itens formatados
            for row in dados:
                preco = f"R$ {row[2]:.2f}" if row[2] is not None else "N/A"
                unidade = row[3] if row[3] else "unidade"
                tree.insert('', 'end', values=(row[0], row[1], preco, unidade))


    # =============================================
    # MÓDULO SUPRIMENTOS
    # =============================================
    
    def criar_suprimentos(self):
        """Cria a interface para gestão de suprimentos"""
        self.frame_suprimentos = Frame(self.frame_conteudo, bg=COR_FUNDO)
        
        # Treeview para listar reposição de suprimentos
        frame_tree = Frame(self.frame_suprimentos, bg=COR_FUNDO)
        frame_tree.pack(fill=BOTH, expand=True, padx=10, pady=10)
        
        scrollbar = Scrollbar(frame_tree)
        scrollbar.pack(side=RIGHT, fill=Y)
        
        self.tree_suprimentos = ttk.Treeview(frame_tree, 
                                           columns=('id', 'imovel', 'item', 'quantidade', 'data', 'valor_gasto', 'comprovante'),
                                           yscrollcommand=scrollbar.set)
        self.tree_suprimentos.pack(fill=BOTH, expand=True)
        scrollbar.config(command=self.tree_suprimentos.yview)
        
        self.tree_suprimentos.heading('#0', text='ID')
        self.tree_suprimentos.heading('#1', text='Imóvel')
        self.tree_suprimentos.heading('#2', text='Item')
        self.tree_suprimentos.heading('#3', text='Quantidade')
        self.tree_suprimentos.heading('#4', text='Data')
        self.tree_suprimentos.heading('#5', text='Valor Gasto')
        self.tree_suprimentos.heading('#6', text='Comprovante')
        
        # Formulário para registrar reposição de suprimentos
        frame_form = Frame(self.frame_suprimentos, bg=COR_CARD, bd=0, 
                         highlightthickness=1, highlightbackground="#e0e0e0")
        frame_form.pack(fill=X, padx=10, pady=10)
        
        Label(frame_form, text="Reposição de Suprimentos", bg=COR_CARD, 
              fg=COR_TEXTO, font=self.fonte_titulo).pack(pady=(10, 5), anchor="w", padx=10)
        
        # Combobox para selecionar imóvel
        frame_imovel = Frame(frame_form, bg=COR_CARD)
        frame_imovel.pack(fill=X, padx=10, pady=5)
        Label(frame_imovel, text="Imóvel:", bg=COR_CARD).pack(side=LEFT, padx=5)
        self.combo_imovel_suprimento = ttk.Combobox(frame_imovel)
        self.combo_imovel_suprimento.pack(side=LEFT, expand=True, fill=X)
        
        # Combobox para selecionar item
        frame_item = Frame(frame_form, bg=COR_CARD)
        frame_item.pack(fill=X, padx=10, pady=5)
        Label(frame_item, text="Item:", bg=COR_CARD).pack(side=LEFT, padx=5)
        self.combo_item_suprimento = ttk.Combobox(frame_item)
        self.combo_item_suprimento.pack(side=LEFT, expand=True, fill=X)
        
        # Campos de quantidade, data e valor
        campos = [
            ("Quantidade:", Entry(frame_form)),
            ("Data:", DateEntry(frame_form, date_pattern='dd/mm/yyyy')),
            ("Valor Gasto (R$):", Entry(frame_form))
        ]
        
        self.entry_suprimento_quantidade = campos[0][1]
        self.entry_suprimento_data = campos[1][1]
        self.entry_suprimento_valor = campos[2][1]
        
        for texto, widget in campos:
            frame = Frame(frame_form, bg=COR_CARD)
            frame.pack(fill=X, padx=10, pady=5)
            Label(frame, text=texto, bg=COR_CARD).pack(side=LEFT, padx=5)
            widget.pack(side=LEFT, expand=True, fill=X)
        
        # Campo para comprovante
        frame_comprovante = Frame(frame_form, bg=COR_CARD)
        frame_comprovante.pack(fill=X, padx=10, pady=5)
        Label(frame_comprovante, text="Comprovante:", bg=COR_CARD).pack(side=LEFT, padx=5)
        self.entry_suprimento_comprovante = Entry(frame_comprovante)
        self.entry_suprimento_comprovante.pack(side=LEFT, expand=True, fill=X)
        Button(frame_comprovante, text="Selecionar", command=lambda: selecionar_arquivo(self.entry_suprimento_comprovante),
              bg=COR_SECUNDARIA, fg="white").pack(side=LEFT, padx=5)
        
        # Botões
        frame_botoes = Frame(frame_form, bg=COR_CARD)
        frame_botoes.pack(fill=X, padx=10, pady=10)
        
        Button(frame_botoes, text="Adicionar", command=self.adicionar_reposicao_suprimento,
              bg=COR_DESTAQUE, fg="white").pack(side=LEFT, padx=5)
        Button(frame_botoes, text="Limpar", command=self.limpar_form_suprimento,
              bg=COR_SECUNDARIA, fg="white").pack(side=LEFT, padx=5)
    
    def carregar_suprimentos(self):
        """Carrega os suprimentos no TreeView e atualiza os comboboxes"""
        conn = sqlite3.connect("sistema.db")
        cursor = conn.cursor()
        
        # Carregar imóveis no combobox
        cursor.execute("SELECT id, endereco FROM imoveis")
        imoveis = cursor.fetchall()
        self.combo_imovel_suprimento['values'] = [f"{i[0]} - {i[1]}" for i in imoveis]
        
        if imoveis:
            self.combo_imovel_suprimento.current(0)
        
        # Carregar itens de suprimento no combobox
        cursor.execute("SELECT id, nome FROM suprimentos")
        itens = cursor.fetchall()
        self.combo_item_suprimento['values'] = [f"{i[0]} - {i[1]}" for i in itens]
        
        if itens:
            self.combo_item_suprimento.current(0)
        
        # Carregar reposição de suprimentos no treeview
        cursor.execute("""
            SELECT r.id, i.endereco, s.nome, r.quantidade, r.data, r.valor_gasto, r.comprovante_path
            FROM reposicao_suprimentos r
            JOIN imoveis i ON r.imovel_id = i.id
            JOIN suprimentos s ON r.suprimento_id = s.id
            ORDER BY r.data DESC
        """)
        
        # Limpar treeview
        for item in self.tree_suprimentos.get_children():
            self.tree_suprimentos.delete(item)
        
        # Adicionar novos itens formatados
        for row in cursor.fetchall():
            comprovante = "Sim" if row[6] else "Não"
            self.tree_suprimentos.insert('', 'end', values=(
                row[0], row[1], row[2], row[3], row[4],
                formatar_moeda(row[5]), 
                comprovante
            ))
        
        conn.close()
    
    def adicionar_reposicao_suprimento(self):
        """Adiciona uma nova reposição de suprimento ao banco de dados"""
        imovel_id = self.combo_imovel_suprimento.get().split(" - ")[0]
        item_id = self.combo_item_suprimento.get().split(" - ")[0]
        quantidade = self.entry_suprimento_quantidade.get()
        data = self.entry_suprimento_data.get_date()
        valor = self.entry_suprimento_valor.get()
        comprovante = self.entry_suprimento_comprovante.get()
        
        if imovel_id and item_id and quantidade and data and valor:
            try:
                quantidade = int(quantidade)
                valor = float(valor)
                
                conn = sqlite3.connect("sistema.db")
                cursor = conn.cursor()
                cursor.execute("""
                    INSERT INTO reposicao_suprimentos 
                    (imovel_id, suprimento_id, quantidade, data, valor_gasto, comprovante_path)
                    VALUES (?, ?, ?, ?, ?, ?)
                """, (imovel_id, item_id, quantidade, data, valor, comprovante))
                conn.commit()
                conn.close()
                
                messagebox.showinfo("Sucesso", "Reposição de suprimento registrada com sucesso!")
                self.carregar_suprimentos()
                self.limpar_form_suprimento()
            except ValueError:
                messagebox.showerror("Erro", "Quantidade deve ser inteiro e valor deve ser numérico")
        else:
            messagebox.showerror("Erro", "Preencha todos os campos obrigatórios")
    
    def limpar_form_suprimento(self):
        """Limpa o formulário de suprimentos"""
        self.entry_suprimento_quantidade.delete(0, END)
        self.entry_suprimento_valor.delete(0, END)
        self.entry_suprimento_comprovante.delete(0, END)
    
    # =============================================
    # MÓDULO RELATÓRIOS
    # =============================================
    
    def criar_relatorios(self):
        """Cria a interface para geração de relatórios profissionais"""
        self.frame_relatorios = Frame(self.frame_conteudo, bg=COR_FUNDO)
        
        # Frame para os botões de relatório
        frame_botoes = Frame(self.frame_relatorios, bg=COR_FUNDO)
        frame_botoes.pack(fill=X, padx=10, pady=20)
        
        # Botões em um frame horizontal
        frame_botoes_horizontal = Frame(frame_botoes, bg=COR_FUNDO)
        frame_botoes_horizontal.pack(fill=X)
        
        # Botão para gerar relatório semanal
        btn_relatorio_semanal = Button(frame_botoes_horizontal, 
                                    text="Gerar Relatório Semanal",
                                    command=self._gerar_relatorio_semanal,
                                    bg=COR_DESTAQUE,
                                    fg="white",
                                    font=("Arial", 10, "bold"),
                                    padx=15,
                                    pady=8)
        btn_relatorio_semanal.pack(side=LEFT, padx=5)
        
        # Botão para gerar relatório mensal
        btn_relatorio_mensal = Button(frame_botoes_horizontal, 
                                    text="Gerar Relatório Mensal",
                                    command=self._gerar_relatorio_mensal,
                                    bg=COR_SECUNDARIA,
                                    fg="white",
                                    font=("Arial", 10, "bold"),
                                    padx=15,
                                    pady=8)
        btn_relatorio_mensal.pack(side=LEFT, padx=5)
        
        # Botão para exportar para PDF
        btn_exportar_pdf = Button(frame_botoes_horizontal, 
                                text="Exportar para PDF",
                                command=self.exportar_para_pdf,
                                bg="#4CAF50",
                                fg="white",
                                font=("Arial", 10, "bold"),
                                padx=15,
                                pady=8)
        btn_exportar_pdf.pack(side=LEFT, padx=5)
        
        # Frame para visualização do relatório
        frame_visualizacao = Frame(self.frame_relatorios, bg=COR_FUNDO)
        frame_visualizacao.pack(fill=BOTH, expand=True, padx=10, pady=10)
        
        # Barra de rolagem
        scrollbar = Scrollbar(frame_visualizacao)
        scrollbar.pack(side=RIGHT, fill=Y)
        
        # Área de texto com barra de rolagem
        self.texto_relatorio = Text(frame_visualizacao, 
                                wrap=WORD, 
                                bg="white",
                                fg="#333333",
                                font=("Arial", 10),
                                yscrollcommand=scrollbar.set,
                                padx=10,
                                pady=10)
        self.texto_relatorio.pack(fill=BOTH, expand=True)
        scrollbar.config(command=self.texto_relatorio.yview)
        
        # Configurar tags para formatação
        self.texto_relatorio.tag_configure("titulo", font=("Arial", 14, "bold"), foreground=COR_PRIMARIA)
        self.texto_relatorio.tag_configure("cabecalho", font=("Arial", 12, "bold"), foreground=COR_SECUNDARIA)
        self.texto_relatorio.tag_configure("subtitulo", font=("Arial", 11, "italic"), foreground="#666666")
        self.texto_relatorio.tag_configure("destaque", font=("Arial", 10, "bold"), foreground=COR_DESTAQUE)
        self.texto_relatorio.tag_configure("negrito", font=("Arial", 10, "bold"))
        self.texto_relatorio.tag_configure("normal", font=("Arial", 10))
        
        # Frame para botões de ação
        frame_acoes = Frame(frame_visualizacao, bg=COR_FUNDO)
        frame_acoes.pack(fill=X, pady=(5, 0))
        
        # Botão para abrir relatório no Word
        btn_abrir_word = Button(frame_acoes,
                            text="Abrir no Word",
                            command=self.abrir_relatorio_word,
                            bg=COR_SECUNDARIA,
                            fg="white",
                            font=("Arial", 9, "bold"),
                            padx=10,
                            pady=5)
        btn_abrir_word.pack(side=RIGHT, padx=5)
        
        # Botão para copiar relatório
        btn_copiar = Button(frame_acoes,
                        text="Copiar Relatório",
                        command=self.copiar_relatorio,
                        bg="#607D8B",
                        fg="white",
                        font=("Arial", 9, "bold"),
                        padx=10,
                        pady=5)
        btn_copiar.pack(side=RIGHT, padx=5)
        
        # Botão para limpar relatório
        btn_limpar = Button(frame_acoes,
                        text="Limpar",
                        command=self.limpar_relatorio,
                        bg=COR_ALERTA,
                        fg="white",
                        font=("Arial", 9, "bold"),
                        padx=10,
                        pady=5)
        btn_limpar.pack(side=RIGHT, padx=5)
    
    def _gerar_relatorio_semanal(self):
        """Gera e exibe o relatório semanal na interface"""
        try:
            caminho = self._gerar_relatorio_semanal_docx()
            
            # Limpar a área de texto
            self.texto_relatorio.delete(1.0, END)
            
            # Adicionar cabeçalho personalizado
            self.adicionar_cabecalho_relatorio("Relatório Semanal")
            
            # Exibir conteúdo do relatório com formatação
            with open(caminho, 'rb') as f:
                doc = Document(f)
                
                for para in doc.paragraphs:
                    texto = para.text
                    estilo = "normal"
                    
                    if para.style.name == 'Heading 1':
                        estilo = "titulo"
                    elif para.style.name == 'Heading 2':
                        estilo = "cabecalho"
                    elif para.style.name == 'Heading 3':
                        estilo = "subtitulo"
                    elif any(run.bold for run in para.runs):
                        estilo = "negrito"
                    
                    self.texto_relatorio.insert(END, texto + "\n", estilo)
            
            messagebox.showinfo("Sucesso", f"Relatório semanal gerado com sucesso!\n{caminho}")
        except Exception as e:
            messagebox.showerror("Erro", f"Erro ao gerar relatório: {str(e)}")

    def _gerar_relatorio_mensal(self):
        """Gera e exibe o relatório mensal na interface"""
        try:
            caminho = self._gerar_relatorio_mensal_docx()
            
            # Limpar a área de texto
            self.texto_relatorio.delete(1.0, END)
            
            # Adicionar cabeçalho personalizado
            self.adicionar_cabecalho_relatorio("Relatório Mensal")
            
            # Exibir conteúdo do relatório com formatação
            with open(caminho, 'rb') as f:
                doc = Document(f)
                
                for para in doc.paragraphs:
                    texto = para.text
                    estilo = "normal"
                    
                    if para.style.name == 'Heading 1':
                        estilo = "titulo"
                    elif para.style.name == 'Heading 2':
                        estilo = "cabecalho"
                    elif para.style.name == 'Heading 3':
                        estilo = "subtitulo"
                    elif any(run.bold for run in para.runs):
                        estilo = "negrito"
                    
                    self.texto_relatorio.insert(END, texto + "\n", estilo)
            
            messagebox.showinfo("Sucesso", f"Relatório mensal gerado com sucesso!\n{caminho}")
        except Exception as e:
            messagebox.showerror("Erro", f"Erro ao gerar relatório mensal: {str(e)}")

       
 
    def adicionar_cabecalho_relatorio(self, titulo):
        """Adiciona um cabeçalho profissional ao relatório"""
        # Informações da empresa
        nome_empresa = "GT - Gestão e serviços para espaços locados por temporada"
        endereco = "Rua dos Imigrantes, 380 - Centro. Pomerode/SC"
        contato = "santos.gustavoethais@gmail.com | (47) 9 9291-2825"
        data_atual = datetime.now().strftime("%d/%m/%Y %H:%M")
        
        # Adicionar cabeçalho formatado
        self.texto_relatorio.insert(END, nome_empresa + "\n", "titulo")
        self.texto_relatorio.insert(END, endereco + "\n", "subtitulo")
        self.texto_relatorio.insert(END, contato + "\n", "subtitulo")
        self.texto_relatorio.insert(END, "\n" + "="*80 + "\n", "normal")
        self.texto_relatorio.insert(END, titulo.upper() + "\n", "cabecalho")
        self.texto_relatorio.insert(END, f"Emitido em: {data_atual}\n", "subtitulo")
        self.texto_relatorio.insert(END, "="*80 + "\n\n", "normal")

    def _gerar_relatorio_semanal_docx(self):
        """Gera um relatório semanal profissional em formato DOCX"""
        conn = sqlite3.connect("sistema.db")
        cursor = conn.cursor()
        
        # Data de início (7 dias atrás)
        data_inicio = (datetime.now() - timedelta(days=7)).strftime("%Y-%m-%d")
        data_fim = datetime.now().strftime("%Y-%m-%d")
        
        # Criar documento profissional
        doc = Document()
        
        # Estilos personalizados
        styles = doc.styles
        style = styles.add_style('CabecalhoEmpresa', WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = 'Arial'
        style.font.size = Pt(14)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)  # COR_PRIMARIA
        
        # Cabeçalho do documento
        para = doc.add_paragraph("SUA EMPRESA LTDA", style='CabecalhoEmpresa')
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        para = doc.add_paragraph("Relatório Semanal de Serviços", style='Heading 1')
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph(f"Período: {data_inicio} a {data_fim}", style='Heading 2')
        
        # Limpezas realizadas
        doc.add_heading('Limpezas Realizadas', level=1)
        cursor.execute("""
            SELECT i.endereco, l.data, l.horas_trabalhadas, l.valor_total 
            FROM limpezas l
            JOIN imoveis i ON l.imovel_id = i.id
            WHERE date(l.data) BETWEEN date(?) AND date(?)
            ORDER BY l.data
        """, (data_inicio, data_fim))
        
        tabela_limpezas = doc.add_table(rows=1, cols=4, style='Light Shading Accent 1')
        tabela_limpezas.autofit = True
        
        # Cabeçalho da tabela
        cabecalho = tabela_limpezas.rows[0].cells
        cabecalho[0].text = 'Imóvel'
        cabecalho[1].text = 'Data'
        cabecalho[2].text = 'Horas'
        cabecalho[3].text = 'Valor'
        
        total_horas = 0
        total_valor = 0
        
        for row in cursor.fetchall():
            linha = tabela_limpezas.add_row().cells
            linha[0].text = row[0]
            linha[1].text = row[1]
            linha[2].text = f"{row[2]:.2f}h"
            linha[3].text = formatar_moeda(row[3])
            total_horas += row[2]
            total_valor += row[3]
        
        doc.add_paragraph(f"Total de horas trabalhadas: {total_horas:.2f}h", style='Body Text')
        doc.add_paragraph(f"Total a receber por limpezas: {formatar_moeda(total_valor)}", style='Body Text')
        
        # Restante da implementação do relatório...
        # ... (manter o restante do código de geração do relatório)

        # Salvar documento
        os.makedirs("relatorios", exist_ok=True)
        nome_arquivo = f"relatorios/Relatorio_Semanal_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
        doc.save(nome_arquivo)
        conn.close()
        
        return nome_arquivo

# Adicione os outros métodos auxiliares (adicionar_cabecalho_relatorio, abrir_relatorio_word, etc.)

    def abrir_relatorio_word(self):
        """Abre o último relatório gerado no Word"""
        try:
            arquivos = [f for f in os.listdir("relatorios") if f.endswith(".docx")]
            if arquivos:
                arquivos.sort(reverse=True)
                caminho = os.path.join("relatorios", arquivos[0])
                os.startfile(caminho)  # Abre com o programa padrão
            else:
                messagebox.showwarning("Aviso", "Nenhum relatório encontrado na pasta 'relatorios'")
        except Exception as e:
            messagebox.showerror("Erro", f"Não foi possível abrir o relatório: {str(e)}")

    def exportar_para_pdf(self):
        """Exporta o relatório atual para PDF"""
        try:
            # Implemente a conversão para PDF aqui
            # Pode usar bibliotecas como python-docx para DOCX para PDF
            messagebox.showinfo("Em desenvolvimento", "Funcionalidade de exportar para PDF será implementada em breve")
        except Exception as e:
            messagebox.showerror("Erro", f"Erro ao exportar para PDF: {str(e)}")

    def copiar_relatorio(self):
        """Copia o conteúdo do relatório para a área de transferência"""
        try:
            conteudo = self.texto_relatorio.get(1.0, END)
            self.root.clipboard_clear()
            self.root.clipboard_append(conteudo)
            messagebox.showinfo("Sucesso", "Relatório copiado para a área de transferência!")
        except Exception as e:
            messagebox.showerror("Erro", f"Erro ao copiar relatório: {str(e)}")

    def limpar_relatorio(self):
        """Limpa a área de visualização do relatório"""
        self.texto_relatorio.delete(1.0, END)

        # =======FECHAMENTO DE CONTAS=========))

    def _construir_aba_fechamento(self, frame, tipo):
        """Constrói a interface de fechamento de contas para imóvel ou cliente"""

        label_titulo = Label(frame, text=f"Fechamento por {tipo.title()}", font=("Arial", 12, "bold"), bg=COR_FUNDO)
        label_titulo.pack(pady=10)

        frame_form = Frame(frame, bg=COR_FUNDO)
        frame_form.pack(fill=X, padx=10, pady=5)

        # Combobox com imóveis ou clientes
        label_item = Label(frame_form, text=f"{'Imóvel' if tipo == 'imovel' else 'Cliente'}:", bg=COR_FUNDO)
        label_item.grid(row=0, column=0, sticky=W)

        combo = ttk.Combobox(frame_form, state="readonly", width=50)
        combo.grid(row=0, column=1, padx=5)
        setattr(self, f"combo_fechamento_{tipo}", combo)

        # Datas
        label_inicio = Label(frame_form, text="Data Início:", bg=COR_FUNDO)
        label_inicio.grid(row=1, column=0, sticky=W)

        entry_inicio = DateEntry(frame_form, width=12, locale='pt_BR')
        entry_inicio.grid(row=1, column=1, sticky=W, padx=5)
        setattr(self, f"entry_data_inicio_{tipo}", entry_inicio)

        label_fim = Label(frame_form, text="Data Fim:", bg=COR_FUNDO)
        label_fim.grid(row=2, column=0, sticky=W)

        entry_fim = DateEntry(frame_form, width=12, locale='pt_BR')
        entry_fim.grid(row=2, column=1, sticky=W, padx=5)
        setattr(self, f"entry_data_fim_{tipo}", entry_fim)

        # Botões
        frame_botoes = Frame(frame, bg=COR_FUNDO)
        frame_botoes.pack(pady=5)

        Button(frame_botoes, text="Gerar Resumo", command=lambda: self.gerar_resumo_fechamento(tipo)).pack(side=LEFT, padx=5)
        Button(frame_botoes, text="Fechar Contas", command=lambda: self.gerar_fechamento(tipo)).pack(side=LEFT, padx=5)
        Button(frame_botoes, text="Limpar", command=lambda: self.limpar_form_fechamento(tipo)).pack(side=LEFT, padx=5)

        # Label de resumo
        label_resumo = Label(frame, text="Selecione um item e o período para ver o resumo", bg=COR_FUNDO, justify=LEFT, anchor=W)
        label_resumo.pack(fill=X, padx=10, pady=10)
        setattr(self, f"label_resumo_{tipo}", label_resumo)

        # Carrega os itens iniciais no combobox
        self.carregar_itens_fechamento(tipo)
    # Adiciona ao menu lateral (no método criar_menu_lateral())
    # Adicione esta linha à lista de módulos:
    # ("Fechar Contas", "fechamento"),

    def _construir_aba_fechamento(self, parent_frame, tipo):
        """Método auxiliar para construir a interface de fechamento"""
        # Frame para seleção
        frame_selecao = Frame(parent_frame, bg=COR_CARD, padx=10, pady=10)
        frame_selecao.pack(fill=X)
        
        Label(frame_selecao, text=f"Selecione o {'Imóvel' if tipo == 'imovel' else 'Cliente'}:", 
            bg=COR_CARD).pack(side=LEFT, padx=5)
        
        setattr(self, f"combo_fechamento_{tipo}", ttk.Combobox(frame_selecao))
        getattr(self, f"combo_fechamento_{tipo}").pack(side=LEFT, expand=True, fill=X, padx=5)
    
        # Frame para período
        frame_periodo = Frame(parent_frame, bg=COR_CARD, padx=10, pady=10)
        frame_periodo.pack(fill=X)
        
        Label(frame_periodo, text="Período:", bg=COR_CARD).grid(row=0, column=0, sticky='e', padx=5)
        Label(frame_periodo, text="De:", bg=COR_CARD).grid(row=0, column=1, sticky='e', padx=5)
        setattr(self, f"entry_data_inicio_{tipo}", DateEntry(frame_periodo, date_pattern='dd/mm/yyyy'))
        getattr(self, f"entry_data_inicio_{tipo}").grid(row=0, column=2, padx=5)
        Label(frame_periodo, text="Até:", bg=COR_CARD).grid(row=0, column=3, sticky='e', padx=5)
        setattr(self, f"entry_data_fim_{tipo}", DateEntry(frame_periodo, date_pattern='dd/mm/yyyy'))
        getattr(self, f"entry_data_fim_{tipo}").grid(row=0, column=4, padx=5)
        
        # Frame para resumo
        frame_resumo = Frame(parent_frame, bg=COR_FUNDO)
        frame_resumo.pack(fill=X, padx=10, pady=5)
        label_resumo = Label(frame_resumo, text="Selecione um item e o período para ver o resumo", bg=COR_FUNDO, justify=LEFT, anchor=W)
        label_resumo.pack(fill=X)
        setattr(self, f"label_resumo_{tipo}", label_resumo)

        # Frame para botões
        frame_botoes = Frame(parent_frame, bg=COR_CARD, padx=10, pady=10)
        frame_botoes.pack(fill=X)
        
        Button(frame_botoes, text="Gerar Resumo", 
            command=lambda: self.gerar_resumo_fechamento(tipo),
            bg=COR_DESTAQUE, fg="white").pack(side=LEFT, padx=5)
        
        Button(frame_botoes, text="Gerar Fechamento", 
            command=lambda: self.gerar_fechamento(tipo),
            bg=COR_SUCESSO, fg="white").pack(side=LEFT, padx=5)
        
        Button(frame_botoes, text="Limpar", 
            command=lambda: self.limpar_form_fechamento(tipo),
            bg=COR_SECUNDARIA, fg="white").pack(side=LEFT, padx=5)
        
        # Carrega os itens no combobox
        self.carregar_itens_fechamento(tipo)

    def carregar_itens_fechamento(self, tipo):
        """Carrega imóveis ou clientes no combobox correspondente ao tipo"""
        conn = sqlite3.connect("sistema.db")
        cursor = conn.cursor()
        
        if tipo == "imovel":
            cursor.execute("SELECT id, endereco FROM imoveis")
            itens = [f"{row[0]} - {row[1]}" for row in cursor.fetchall()]
        else:
            cursor.execute("SELECT id, nome FROM clientes")
            itens = [f"{row[0]} - {row[1]}" for row in cursor.fetchall()]
        
        combo = getattr(self, f"combo_fechamento_{tipo}", None)
        if combo:
            combo['values'] = itens
            if itens:
                combo.current(0)
        
        conn.close()


    def gerar_resumo_fechamento(self, tipo):
        """Gera um resumo dos valores a receber"""
        try:
            combo = getattr(self, f"combo_fechamento_{tipo}")
            entry_inicio = getattr(self, f"entry_data_inicio_{tipo}")
            entry_fim = getattr(self, f"entry_data_fim_{tipo}")
            label_resumo = getattr(self, f"label_resumo_{tipo}")

            item = combo.get()
            if not item:
                raise ValueError("Selecione um item")
                    
            item_id = item.split(" - ")[0]
            data_inicio = entry_inicio.get_date()
            data_fim = entry_fim.get_date()
                        
            if data_inicio > data_fim:
                    raise ValueError("Data inicial deve ser anterior à data final")
                    
            conn = sqlite3.connect("sistema.db")
            cursor = conn.cursor()
                
                # Calcula totais
            if tipo == "imovel":
                    # Limpezas
                    cursor.execute("""
                        SELECT COALESCE(SUM(valor_total), 0) 
                        FROM limpezas 
                        WHERE imovel_id = ? AND date(data) BETWEEN ? AND ?
                    """, (item_id, data_inicio, data_fim))
                    total_limpezas = cursor.fetchone()[0]
                    
                    # Enxoval
                    cursor.execute("""
                        SELECT COALESCE(SUM(t.preco_unitario * c.quantidade), 0)
                        FROM consumo_enxoval c
                        JOIN tipos_enxoval t ON c.item_id = t.id
                        WHERE c.imovel_id = ? AND date(c.data) BETWEEN ? AND ?
                    """, (item_id, data_inicio, data_fim))
                    total_enxoval = cursor.fetchone()[0]
                    
                    # Suprimentos
                    cursor.execute("""
                        SELECT COALESCE(SUM(valor_gasto), 0)
                        FROM reposicao_suprimentos
                        WHERE imovel_id = ? AND date(data) BETWEEN ? AND ?
                    """, (item_id, data_inicio, data_fim))
                    total_suprimentos = cursor.fetchone()[0]
                    
                    # Valor fixo por gestão (considerando 1 imóvel)
                    valor_gestao = 50.0
                    
                    resumo = f"""
                        RESUMO PARA O IMÓVEL: {item}
                        Período: {data_inicio.strftime('%d/%m/%Y')} a {data_fim.strftime('%d/%m/%Y')}

                        - Limpezas: {formatar_moeda(total_limpezas)}
                        - Enxoval: {formatar_moeda(total_enxoval)}
                        - Suprimentos: {formatar_moeda(total_suprimentos)}
                        - Gestão: {formatar_moeda(valor_gestao)}
                        ---------------------------
                        TOTAL: {formatar_moeda(total_limpezas + total_enxoval + total_suprimentos + valor_gestao)}
                        """
            else:
                    # Para cliente, somamos todos os imóveis dele
                    cursor.execute("SELECT id FROM imoveis WHERE cliente_id = ?", (item_id,))
                    imoveis_ids = [row[0] for row in cursor.fetchall()]
                    
                    if not imoveis_ids:
                        raise ValueError("Este cliente não possui imóveis cadastrados")
                        
                    # Limpezas
                    cursor.execute(f"""
                        SELECT COALESCE(SUM(valor_total), 0) 
                        FROM limpezas 
                        WHERE imovel_id IN ({','.join(['?']*len(imoveis_ids))}) 
                        AND date(data) BETWEEN ? AND ?
                    """, (*imoveis_ids, data_inicio, data_fim))
                    total_limpezas = cursor.fetchone()[0]
                    
                    # Enxoval
                    cursor.execute(f"""
                        SELECT COALESCE(SUM(t.preco_unitario * c.quantidade), 0)
                        FROM consumo_enxoval c
                        JOIN tipos_enxoval t ON c.item_id = t.id
                        WHERE c.imovel_id IN ({','.join(['?']*len(imoveis_ids))})
                        AND date(c.data) BETWEEN ? AND ?
                    """, (*imoveis_ids, data_inicio, data_fim))
                    total_enxoval = cursor.fetchone()[0]
                    
                    # Suprimentos
                    cursor.execute(f"""
                        SELECT COALESCE(SUM(valor_gasto), 0)
                        FROM reposicao_suprimentos
                        WHERE imovel_id IN ({','.join(['?']*len(imoveis_ids))})
                        AND date(data) BETWEEN ? AND ?
                    """, (*imoveis_ids, data_inicio, data_fim))
                    total_suprimentos = cursor.fetchone()[0]
                    
                    # Valor fixo por gestão (por imóvel)
                    valor_gestao = 50.0 * len(imoveis_ids)
                    
                    resumo = f"""
                    RESUMO PARA O CLIENTE: {item}
                    Período: {data_inicio.strftime('%d/%m/%Y')} a {data_fim.strftime('%d/%m/%Y')}
                    Imóveis: {len(imoveis_ids)}

                    - Limpezas: {formatar_moeda(total_limpezas)}
                    - Enxoval: {formatar_moeda(total_enxoval)}
                    - Suprimentos: {formatar_moeda(total_suprimentos)}
                    - Gestão: {formatar_moeda(valor_gestao)}
                    ---------------------------
                    TOTAL: {formatar_moeda(total_limpezas + total_enxoval + total_suprimentos + valor_gestao)}
                    """
                
            self.label_resumo.config(text=resumo)
            conn.close()
                
        except ValueError as e:
                messagebox.showerror("Erro", str(e))
        except Exception as e:
                messagebox.showerror("Erro", f"Ocorreu um erro ao gerar o resumo:\n{str(e)}")     


    def gerar_fechamento(self, tipo):
        """Registra o fechamento no banco de dados"""
        try:
            # Widgets dinâmicos conforme o tipo (imovel ou cliente)
            combo = getattr(self, f"combo_fechamento_{tipo}")
            entry_inicio = getattr(self, f"entry_data_inicio_{tipo}")
            entry_fim = getattr(self, f"entry_data_fim_{tipo}")
            label_resumo = getattr(self, f"label_resumo_{tipo}")

            item = combo.get()
            if not item:
                raise ValueError("Selecione um item")
            
            item_id = item.split(" - ")[0]
            data_inicio = entry_inicio.get_date()
            data_fim = entry_fim.get_date()
            
            if data_inicio > data_fim:
                raise ValueError("Data inicial deve ser anterior à data final")
            
            # Confirmação do usuário
            confirmacao = messagebox.askyesno(
                "Confirmar Fechamento",
                f"Tem certeza que deseja fechar as contas deste {'imóvel' if tipo == 'imovel' else 'cliente'} "
                f"no período de {data_inicio.strftime('%d/%m/%Y')} a {data_fim.strftime('%d/%m/%Y')}?"
            )
            if not confirmacao:
                return

            conn = sqlite3.connect("sistema.db")
            cursor = conn.cursor()

            if tipo == "imovel":
                # Limpezas
                cursor.execute("""
                    SELECT COALESCE(SUM(valor_total), 0) 
                    FROM limpezas 
                    WHERE imovel_id = ? AND date(data) BETWEEN ? AND ?
                """, (item_id, data_inicio, data_fim))
                total_limpezas = cursor.fetchone()[0]

                # Enxoval
                cursor.execute("""
                    SELECT COALESCE(SUM(t.preco_unitario * c.quantidade), 0)
                    FROM consumo_enxoval c
                    JOIN tipos_enxoval t ON c.item_id = t.id
                    WHERE c.imovel_id = ? AND date(c.data) BETWEEN ? AND ?
                """, (item_id, data_inicio, data_fim))
                total_enxoval = cursor.fetchone()[0]

                # Suprimentos
                cursor.execute("""
                    SELECT COALESCE(SUM(valor_gasto), 0)
                    FROM reposicao_suprimentos
                    WHERE imovel_id = ? AND date(data) BETWEEN ? AND ?
                """, (item_id, data_inicio, data_fim))
                total_suprimentos = cursor.fetchone()[0]

                valor_gestao = 50.0
                valor_total = total_limpezas + total_enxoval + total_suprimentos + valor_gestao

            else:  # cliente
                cursor.execute("SELECT id FROM imoveis WHERE cliente_id = ?", (item_id,))
                imoveis_ids = [row[0] for row in cursor.fetchall()]

                if not imoveis_ids:
                    raise ValueError("Este cliente não possui imóveis cadastrados")

                q_marks = ','.join(['?'] * len(imoveis_ids))

                # Limpezas
                cursor.execute(f"""
                    SELECT COALESCE(SUM(valor_total), 0) 
                    FROM limpezas 
                    WHERE imovel_id IN ({q_marks}) AND date(data) BETWEEN ? AND ?
                """, (*imoveis_ids, data_inicio, data_fim))
                total_limpezas = cursor.fetchone()[0]

                # Enxoval
                cursor.execute(f"""
                    SELECT COALESCE(SUM(t.preco_unitario * c.quantidade), 0)
                    FROM consumo_enxoval c
                    JOIN tipos_enxoval t ON c.item_id = t.id
                    WHERE c.imovel_id IN ({q_marks}) AND date(c.data) BETWEEN ? AND ?
                """, (*imoveis_ids, data_inicio, data_fim))
                total_enxoval = cursor.fetchone()[0]

                # Suprimentos
                cursor.execute(f"""
                    SELECT COALESCE(SUM(valor_gasto), 0)
                    FROM reposicao_suprimentos
                    WHERE imovel_id IN ({q_marks}) AND date(data) BETWEEN ? AND ?
                """, (*imoveis_ids, data_inicio, data_fim))
                total_suprimentos = cursor.fetchone()[0]

                valor_gestao = 50.0 * len(imoveis_ids)
                valor_total = total_limpezas + total_enxoval + total_suprimentos + valor_gestao

            # Pergunta sobre comprovante
            comprovante = None
            if messagebox.askyesno("Comprovante", "Deseja anexar um comprovante de pagamento?"):
                caminho = filedialog.askopenfilename()
                if caminho:
                    nome_arquivo = os.path.basename(caminho)
                    destino = os.path.join("comprovantes", nome_arquivo)
                    os.makedirs("comprovantes", exist_ok=True)  # garante que a pasta existe
                    os.replace(caminho, destino)
                    comprovante = destino

            # Salva o fechamento
            cursor.execute("""
                INSERT INTO fechamentos 
                (tipo, referencia_id, data_inicio, data_fim, valor_total, comprovante_path)
                VALUES (?, ?, ?, ?, ?, ?)
            """, (tipo, item_id, data_inicio, data_fim, valor_total, comprovante))

            conn.commit()
            conn.close()

            messagebox.showinfo("Sucesso", f"Fechamento registrado com sucesso!\nTotal: {formatar_moeda(valor_total)}")
            self.limpar_form_fechamento(tipo)

        except ValueError as e:
            messagebox.showerror("Erro", str(e))
        except Exception as e:
            messagebox.showerror("Erro", f"Ocorreu um erro ao registrar o fechamento:\n{str(e)}")

            

    def limpar_form_fechamento(self, tipo):
        """Limpa o formulário de fechamento (cliente ou imóvel)"""

        # Acessa os widgets dinamicamente com base no tipo
        combo = getattr(self, f"combo_fechamento_{tipo}")
        entry_inicio = getattr(self, f"entry_data_inicio_{tipo}")
        entry_fim = getattr(self, f"entry_data_fim_{tipo}")
        label_resumo = getattr(self, f"label_resumo_{tipo}")

        # Limpa o combobox
        combo.set('')

        # Define datas padrão: últimos 30 dias
        hoje = datetime.now().date()
        entry_inicio.set_date(hoje - timedelta(days=30))
        entry_fim.set_date(hoje)

        # Limpa o texto do resumo
        label_resumo.config(text="Selecione um item e o período para ver o resumo")



    def criar_historico_fechamentos(self):
        """Cria a interface para visualizar fechamentos anteriores"""
        self.frame_historico_fechamentos = Frame(self.frame_conteudo, bg=COR_FUNDO)
        
        # Treeview para listar fechamentos
        frame_tree = Frame(self.frame_historico_fechamentos, bg=COR_FUNDO)
        frame_tree.pack(fill=BOTH, expand=True, padx=10, pady=10)
        
        scrollbar = Scrollbar(frame_tree)
        scrollbar.pack(side=RIGHT, fill=Y)
        
        self.tree_fechamentos = ttk.Treeview(frame_tree, 
                                        columns=('id', 'tipo', 'referencia', 'periodo', 'valor', 'data_fechamento'),
                                        yscrollcommand=scrollbar.set)
        self.tree_fechamentos.pack(fill=BOTH, expand=True)
        scrollbar.config(command=self.tree_fechamentos.yview)
        
        self.tree_fechamentos.heading('#0', text='ID')
        self.tree_fechamentos.heading('#1', text='Tipo')
        self.tree_fechamentos.heading('#2', text='Referência')
        self.tree_fechamentos.heading('#3', text='Período')
        self.tree_fechamentos.heading('#4', text='Valor')
        self.tree_fechamentos.heading('#5', text='Data Fechamento')
        
        # Botão para atualizar
        frame_botoes = Frame(self.frame_historico_fechamentos, bg=COR_FUNDO, padx=10, pady=10)
        frame_botoes.pack(fill=X)
        
        Button(frame_botoes, text="Atualizar", command=self.carregar_fechamentos,
            bg=COR_DESTAQUE, fg="white").pack(side=LEFT, padx=5)
        
        # Carrega os fechamentos
        self.carregar_fechamentos()


    def carregar_fechamentos(self):
        """Carrega os fechamentos no TreeView"""
        conn = sqlite3.connect("sistema.db")
        cursor = conn.cursor()
        
        cursor.execute("""
            SELECT f.id, f.tipo, 
                CASE 
                    WHEN f.tipo = 'imovel' THEN i.endereco
                    ELSE c.nome
                END as referencia,
                f.data_inicio || ' a ' || f.data_fim as periodo,
                f.valor_total, f.data_fechamento
            FROM fechamentos f
            LEFT JOIN imoveis i ON f.tipo = 'imovel' AND f.referencia_id = i.id
            LEFT JOIN clientes c ON f.tipo = 'cliente' AND f.referencia_id = c.id
            ORDER BY f.data_fechamento DESC
        """)
        
        # Limpar treeview
        for item in self.tree_fechamentos.get_children():
            self.tree_fechamentos.delete(item)
        
        # Adicionar novos itens formatados
        for row in cursor.fetchall():
            self.tree_fechamentos.insert('', 'end', values=(
                row[0], 
                "Imóvel" if row[1] == "imovel" else "Cliente",
                row[2], 
                row[3], 
                formatar_moeda(row[4]), 
                row[5]
            ))
        
        conn.close()


    def criar_fechamento_contas(self):
        """Cria a interface para o módulo Fechamento de Contas"""
        self.frame_fechamento = Frame(self.frame_conteudo, bg=COR_FUNDO)
        abas = ttk.Notebook(self.frame_fechamento)
        frame_imovel = Frame(abas, bg=COR_FUNDO)
        frame_cliente = Frame(abas, bg=COR_FUNDO)
        abas.add(frame_imovel, text="Por Imóvel")
        abas.add(frame_cliente, text="Por Cliente")
        abas.pack(fill=BOTH, expand=True)
        # Constrói as abas usando seu método auxiliar
        self._construir_aba_fechamento(frame_imovel, "imovel")
        self._construir_aba_fechamento(frame_cliente, "cliente")



# =============================================
# INICIALIZAÇÃO DO SISTEMA
# =============================================

if __name__ == "__main__":
    from datetime import timedelta  # Import necessário para o timedelta
    
    root = Tk()
    app = SistemaGestaoApp(root)
    root.mainloop()
